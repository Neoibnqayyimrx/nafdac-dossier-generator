"""
nafdac_validator.py — Phase 7
7.1  NAFDAC checklist validator  — checks all Modules 1-5 against NAFDAC CTD requirements
7.2  Missing section reporter    — terminal + DOCX report of what's missing/incomplete
7.3  End-to-end integration test — runs the full pipeline and reports status
"""

from __future__ import annotations

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═════════════════════════════════════════════════════════════════════════════
# 7.1  NAFDAC CTD CHECKLIST DEFINITION
# All items are keyed by their CTD section number.
# Each item has:
#   path        — relative path inside the submission folder
#   required    — True = mandatory for NAFDAC new product registration
#   description — human-readable label
#   check       — "exists" | "non_empty" | "min_size" (bytes)
# ═════════════════════════════════════════════════════════════════════════════

NAFDAC_CHECKLIST = {
    # ── Module 1 ─────────────────────────────────────────────────────────────
    "1.0":   {"path": "Module1",                                              "required": True,  "description": "Module 1 folder present",                        "check": "exists"},
    "1.1":   {"path": "Module1/1.1_Cover_Letter",                             "required": True,  "description": "Cover Letter folder",                            "check": "exists"},
    "1.1.1": {"path": "Module1/1.1_Cover_Letter/1.1_cover_letter.docx",           "required": True,  "description": "Cover Letter document",                          "check": "non_empty"},
    "1.2":   {"path": "Module1/1.2_Administrative",                           "required": True,  "description": "Administrative documents folder",                "check": "exists"},
    "1.2.1": {"path": "Module1/1.2_Administrative/m1_application_form.docx",     "required": False, "description": "NAFDAC Application Form (manual — Form 5A)",                        "check": "non_empty"},
    "1.2.2": {"path": "Module1/1.2_Administrative/certificate_of_analysis.docx","required": False,"description": "Certificate of Analysis",                       "check": "exists"},
    "1.3":   {"path": "Module1/1.3_Product_Information",                      "required": True,  "description": "Product Information folder",                     "check": "exists"},
    "1.3.1": {"path": "Module1/1.3_Product_Information/1.3.1_SmPC/smpc.docx", "required": True,  "description": "Summary of Product Characteristics (SmPC)",      "check": "non_empty"},
    "1.3.2": {"path": "Module1/1.3_Product_Information/1.3.2_PIL/pil.docx",   "required": True,  "description": "Patient Information Leaflet (PIL)",              "check": "non_empty"},
    "1.3.3": {"path": "Module1/1.3_Product_Information/1.3.3_Labelling",      "required": True,  "description": "Labelling folder",                               "check": "exists"},
    "1.4":   {"path": "Module1/1.4_GMP_Documents",                            "required": True,  "description": "GMP Documents folder",                           "check": "exists"},
    "1.5":   {"path": "Module1/1.5_References",                               "required": False, "description": "References folder",                              "check": "exists"},

    # ── Module 2 ─────────────────────────────────────────────────────────────
    "2.0":   {"path": "Module2",                                              "required": True,  "description": "Module 2 folder present",                        "check": "exists"},
    "2.2":   {"path": "Module2/2.2_Introduction",                             "required": True,  "description": "Introduction to the dossier",                    "check": "exists"},
    "2.3":   {"path": "Module2/2.3_QOS",                                      "required": True,  "description": "Quality Overall Summary (QOS) folder",           "check": "exists"},
    "2.3.1": {"path": "Module2/2.3_QOS/2.3_quality_overall_summary.docx",         "required": True,  "description": "Quality Overall Summary document",               "check": "non_empty"},
    "2.4":   {"path": "Module2/2.4_Non_Clinical_Overview",                    "required": False, "description": "Non-clinical Overview",                          "check": "exists"},
    "2.5":   {"path": "Module2/2.5_Clinical_Overview",                        "required": False, "description": "Clinical Overview",                              "check": "exists"},
    "2.6":   {"path": "Module2/2.6_Non_Clinical_Summaries",                   "required": False, "description": "Non-clinical Written and Tabulated Summaries",   "check": "exists"},
    "2.7":   {"path": "Module2/2.7_Clinical_Summaries",                       "required": False, "description": "Clinical Summary",                               "check": "exists"},

    # ── Module 3 ─────────────────────────────────────────────────────────────
    "3.0":   {"path": "Module3",                                              "required": True,  "description": "Module 3 folder present",                        "check": "exists"},
    "3.2.S": {"path": "Module3/3.2_Body_of_Data/3.2.S_Drug_Substance",                         "required": True,  "description": "Drug Substance section",                         "check": "exists"},
    "3.2.S.1": {"path": "Module3/3.2_Body_of_Data/3.2.S_Drug_Substance/3.2.S.1_General_Information/3.2.S.1_general_information.docx",
                                                                               "required": True,  "description": "Drug Substance — General Information",           "check": "non_empty"},
    "3.2.S.2": {"path": "Module3/3.2_Body_of_Data/3.2.S_Drug_Substance/3.2.S.2_Manufacture",   "required": True,  "description": "Drug Substance — Manufacture",                   "check": "exists"},
    "3.2.S.3": {"path": "Module3/3.2_Body_of_Data/3.2.S_Drug_Substance/3.2.S.3_Characterisation","required": True, "description": "Drug Substance — Characterisation",              "check": "exists"},
    "3.2.S.4": {"path": "Module3/3.2_Body_of_Data/3.2.S_Drug_Substance/3.2.S.4_Control/3.2.S.4.1_specifications.docx",
                                                                               "required": True,  "description": "Drug Substance — Specifications",                "check": "non_empty"},
    "3.2.S.5": {"path": "Module3/3.2_Body_of_Data/3.2.S_Drug_Substance/3.2.S.5_Reference_Standards","required": True,"description": "Drug Substance — Reference Standards",         "check": "exists"},
    "3.2.S.6": {"path": "Module3/3.2_Body_of_Data/3.2.S_Drug_Substance/3.2.S.6_Container_Closure","required": True,"description": "Drug Substance — Container Closure System",      "check": "exists"},
    "3.2.S.7": {"path": "Module3/3.2_Body_of_Data/3.2.S_Drug_Substance/3.2.S.7_Stability",     "required": True,  "description": "Drug Substance — Stability",                     "check": "exists"},
    "3.2.P": {"path": "Module3/3.2_Body_of_Data/3.2.P_Drug_Product",                           "required": True,  "description": "Drug Product section",                           "check": "exists"},
    "3.2.P.1": {"path": "Module3/3.2_Body_of_Data/3.2.P_Drug_Product/3.2.P.1_Description/3.2.P.1_description_composition.docx",
                                                                               "required": True,  "description": "Drug Product — Description & Composition",       "check": "non_empty"},
    "3.2.P.2": {"path": "Module3/3.2_Body_of_Data/3.2.P_Drug_Product/3.2.P.2_Pharmaceutical_Development","required": True,"description": "Drug Product — Pharmaceutical Development","check": "exists"},
    "3.2.P.3": {"path": "Module3/3.2_Body_of_Data/3.2.P_Drug_Product/3.2.P.3_Manufacture",     "required": True,  "description": "Drug Product — Manufacture",                     "check": "exists"},
    "3.2.P.4": {"path": "Module3/3.2_Body_of_Data/3.2.P_Drug_Product/3.2.P.4_Control_of_Excipients","required": True,"description": "Drug Product — Control of Excipients",            "check": "exists"},
    "3.2.P.5": {"path": "Module3/3.2_Body_of_Data/3.2.P_Drug_Product/3.2.P.5_Control_of_Drug_Product",
                                                                               "required": True,  "description": "Drug Product — Specifications",                  "check": "exists"},
    "3.2.P.6": {"path": "Module3/3.2_Body_of_Data/3.2.P_Drug_Product/3.2.P.6_Reference_Standards","required": True,"description": "Drug Product — Reference Standards",             "check": "exists"},
    "3.2.P.7": {"path": "Module3/3.2_Body_of_Data/3.2.P_Drug_Product/3.2.P.7_Container_Closure","required": True, "description": "Drug Product — Container Closure System",         "check": "exists"},
    "3.2.P.8": {"path": "Module3/3.2_Body_of_Data/3.2.P_Drug_Product/3.2.P.8_Stability",       "required": True,  "description": "Drug Product — Stability",                       "check": "exists"},
    "3.3":   {"path": "Module3/3.3_Literature_References",                    "required": False, "description": "Literature References",                          "check": "exists"},

    # ── Module 4 ─────────────────────────────────────────────────────────────
    "4.0":   {"path": "Module4",                                              "required": False, "description": "Module 4 folder (Non-clinical Study Reports)",    "check": "exists"},
    "4.2.1": {"path": "Module4/4.2.1_Pharmacology",                           "required": False, "description": "Pharmacology studies",                           "check": "exists"},
    "4.2.2": {"path": "Module4/4.2.2_Pharmacokinetics",                       "required": False, "description": "Pharmacokinetics studies",                       "check": "exists"},
    "4.2.3": {"path": "Module4/4.2.3_Toxicology",                             "required": False, "description": "Toxicology studies",                             "check": "exists"},

    # ── Module 5 ─────────────────────────────────────────────────────────────
    "5.0":   {"path": "Module5",                                              "required": False, "description": "Module 5 folder (Clinical Study Reports)",        "check": "exists"},
    "5.2":   {"path": "Module5/5.2_Tabular_Listing",                          "required": False, "description": "Tabular Listing of Clinical Studies",            "check": "exists"},
    "5.3":   {"path": "Module5/5.3_Clinical_Study_Reports",                   "required": False, "description": "Clinical Study Reports",                         "check": "exists"},
}

# Status constants
STATUS_PASS    = "PASS"
STATUS_MISSING = "MISSING"
STATUS_EMPTY   = "EMPTY"
STATUS_WARN    = "WARNING"
STATUS_NA      = "N/A"


# ═════════════════════════════════════════════════════════════════════════════
# 7.1  VALIDATOR CORE
# ═════════════════════════════════════════════════════════════════════════════

def validate_submission(submission_root: str | Path) -> dict:
    """
    Validate a submission folder against the NAFDAC CTD checklist.

    Returns:
        {
          "submission_root": str,
          "product_name": str,
          "validated_at": str,
          "results": { section_id: { status, description, path, required, note } },
          "summary": {
            "total": int, "passed": int, "missing_required": int,
            "missing_optional": int, "empty": int, "warnings": int,
            "score": float,   # % of required items passing
            "ready_to_submit": bool,
          },
          "missing_required": [ { id, description, path } ],
          "missing_optional": [ { id, description, path } ],
          "warnings":         [ str ],
        }
    """
    root = Path(submission_root).resolve()
    results = {}
    missing_required = []
    missing_optional = []
    warnings_list = []

    # Load manifest for product name
    product_name = root.name
    manifest_path = root / "submission_manifest.json"
    if manifest_path.exists():
        try:
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            product_name = manifest.get("full_product_name", product_name)
        except Exception:
            pass

    for section_id, item in NAFDAC_CHECKLIST.items():
        item_path = root / item["path"]
        required  = item["required"]
        desc      = item["description"]
        check     = item["check"]

        entry = {
            "status":      STATUS_NA,
            "description": desc,
            "path":        item["path"],
            "required":    required,
            "note":        "",
        }

        if not item_path.exists():
            if required:
                entry["status"] = STATUS_MISSING
                entry["note"]   = "Required — not found"
                missing_required.append({"id": section_id, "description": desc, "path": item["path"]})
            else:
                entry["status"] = STATUS_NA
                entry["note"]   = "Optional — not present"
                missing_optional.append({"id": section_id, "description": desc, "path": item["path"]})

        elif check == "non_empty":
            size = item_path.stat().st_size if item_path.is_file() else 0
            if size < 1000:  # under 1 KB = likely a stub/placeholder
                entry["status"] = STATUS_EMPTY
                entry["note"]   = f"File exists but appears empty or stub ({size} bytes)"
                if required:
                    warnings_list.append(f"{section_id} — {desc}: file is very small ({size} bytes), may be a stub")
            else:
                entry["status"] = STATUS_PASS
                entry["note"]   = f"{size:,} bytes"

        elif check == "min_size":
            min_b = item.get("min_bytes", 500)
            size  = item_path.stat().st_size if item_path.is_file() else 0
            if size < min_b:
                entry["status"] = STATUS_WARN
                entry["note"]   = f"File smaller than expected ({size} bytes, min {min_b})"
                warnings_list.append(f"{section_id} — {desc}: file may be incomplete")
            else:
                entry["status"] = STATUS_PASS

        else:  # "exists"
            entry["status"] = STATUS_PASS
            if item_path.is_dir():
                n_files = len(list(item_path.rglob("*")))
                entry["note"] = f"{n_files} item(s) inside"
            else:
                entry["note"] = "Present"

        results[section_id] = entry

    # ── Summary stats ─────────────────────────────────────────────────────
    total            = len(NAFDAC_CHECKLIST)
    passed           = sum(1 for r in results.values() if r["status"] == STATUS_PASS)
    empty_count      = sum(1 for r in results.values() if r["status"] == STATUS_EMPTY)
    miss_req_count   = len(missing_required)
    miss_opt_count   = len(missing_optional)
    warn_count       = len(warnings_list)

    required_total   = sum(1 for i in NAFDAC_CHECKLIST.values() if i["required"])
    required_passing = sum(
        1 for sid, r in results.items()
        if NAFDAC_CHECKLIST[sid]["required"] and r["status"] == STATUS_PASS
    )
    score = round((required_passing / required_total) * 100, 1) if required_total else 0.0
    ready = miss_req_count == 0 and empty_count == 0

    return {
        "submission_root":  str(root),
        "product_name":     product_name,
        "validated_at":     datetime.now().strftime("%Y-%m-%d %H:%M"),
        "results":          results,
        "summary": {
            "total":              total,
            "passed":             passed,
            "missing_required":   miss_req_count,
            "missing_optional":   miss_opt_count,
            "empty":              empty_count,
            "warnings":           warn_count,
            "score":              score,
            "ready_to_submit":    ready,
        },
        "missing_required": missing_required,
        "missing_optional": missing_optional,
        "warnings":         warnings_list,
    }


# ═════════════════════════════════════════════════════════════════════════════
# 7.2  DOCX REPORT GENERATOR
# ═════════════════════════════════════════════════════════════════════════════

def _cell_shading(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _status_color(status: str) -> tuple[str, str]:
    """Returns (hex_bg, hex_text) for each status."""
    return {
        STATUS_PASS:    ("E2EFDA", "375A23"),
        STATUS_MISSING: ("FCE4D6", "C00000"),
        STATUS_EMPTY:   ("FFF2CC", "7F6000"),
        STATUS_WARN:    ("FFF2CC", "7F6000"),
        STATUS_NA:      ("F2F2F2", "595959"),
    }.get(status, ("FFFFFF", "000000"))


def save_validation_report(
    validation_result: dict,
    output_path: Path | str,
) -> Path:
    """
    Save the full validation result as a formatted DOCX report.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc  = Document()
    summ = validation_result["summary"]

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_heading("NAFDAC Dossier Validation Report", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    meta = doc.add_paragraph()
    meta.add_run("Product: ").bold = True
    meta.add_run(validation_result["product_name"])
    meta.add_run("    Validated: ").bold = True
    meta.add_run(validation_result["validated_at"])
    meta.add_run("    Root: ").bold = True
    meta.add_run(str(validation_result["submission_root"]))

    doc.add_paragraph()

    # ── Score banner ───────────────────────────────────────────────────────
    score_para = doc.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_run = score_para.add_run(
        f"Compliance Score: {summ['score']}%  |  "
        f"{'✓ READY TO SUBMIT' if summ['ready_to_submit'] else '✗ NOT READY — SEE BELOW'}"
    )
    score_run.bold = True
    score_run.font.size = Pt(14)
    score_run.font.color.rgb = (
        RGBColor(0x37, 0x5A, 0x23) if summ["ready_to_submit"]
        else RGBColor(0xC0, 0x00, 0x00)
    )

    doc.add_paragraph()

    # ── Summary table ──────────────────────────────────────────────────────
    doc.add_heading("Summary", level=2)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Count"
    for h in hdr:
        h.paragraphs[0].runs[0].bold = True
        _cell_shading(h, "1F497D")
        h.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for label, val in [
        ("Total checklist items",         summ["total"]),
        ("✓ Passed",                       summ["passed"]),
        ("✗ Missing (required)",           summ["missing_required"]),
        ("⚠ Empty / stub files",           summ["empty"]),
        ("– Missing (optional)",           summ["missing_optional"]),
        ("⚠ Warnings",                     summ["warnings"]),
    ]:
        row = summary_table.add_row().cells
        row[0].text = label
        row[1].text = str(val)

    doc.add_paragraph()

    # ── Missing required items ─────────────────────────────────────────────
    if validation_result["missing_required"]:
        doc.add_heading("Missing Required Items", level=2)
        mr_table = doc.add_table(rows=1, cols=3)
        mr_table.style = "Table Grid"
        for i, h in enumerate(["Section", "Description", "Expected Path"]):
            cell = mr_table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _cell_shading(cell, "C00000")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for item in validation_result["missing_required"]:
            row = mr_table.add_row().cells
            row[0].text = item["id"]
            row[1].text = item["description"]
            row[2].text = item["path"]
            for cell in row:
                _cell_shading(cell, "FCE4D6")
        doc.add_paragraph()

    # ── Warnings ───────────────────────────────────────────────────────────
    if validation_result["warnings"]:
        doc.add_heading("Warnings", level=2)
        for w in validation_result["warnings"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(w).font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
        doc.add_paragraph()

    # ── Full checklist ─────────────────────────────────────────────────────
    doc.add_heading("Full Checklist", level=2)
    cl_table = doc.add_table(rows=1, cols=5)
    cl_table.style = "Table Grid"

    for i, h in enumerate(["ID", "Description", "Required", "Status", "Note"]):
        cell = cl_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        _cell_shading(cell, "1F497D")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for section_id, item in validation_result["results"].items():
        row   = cl_table.add_row().cells
        bg, _ = _status_color(item["status"])
        row[0].text = section_id
        row[1].text = item["description"]
        row[2].text = "Yes" if item["required"] else "No"
        row[3].text = item["status"]
        row[4].text = item.get("note", "")
        for cell in row:
            _cell_shading(cell, bg)

    doc.add_paragraph()

    # ── Optional missing ───────────────────────────────────────────────────
    if validation_result["missing_optional"]:
        doc.add_heading("Optional Sections Not Present", level=2)
        doc.add_paragraph(
            "These sections are not mandatory for NAFDAC new product registration "
            "but may be requested during review.",
            style="Normal",
        )
        for item in validation_result["missing_optional"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{item['id']} — {item['description']}").font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.save(output_path)
    return output_path


# ═════════════════════════════════════════════════════════════════════════════
# 7.3  END-TO-END INTEGRATION TEST
# ═════════════════════════════════════════════════════════════════════════════

def run_integration_test(
    submission_root: str | Path,
    drug_name: str,
    verbose: bool = True,
) -> dict:
    """
    Run the full pipeline check for a submission:
      1. Folder structure exists
      2. Manifest is valid
      3. All required DOCX files present and non-empty
      4. Spec resolved (pharmacopoeia_db/resolved/<drug>.json exists)
      5. SmPC and PIL present
      6. NAFDAC checklist validation
      7. OpenFDA data available
      8. Scraper cache populated

    Returns a dict with per-step results and overall pass/fail.
    """
    root = Path(submission_root).resolve()
    steps = []
    all_passed = True

    def _step(name: str, passed: bool, detail: str = "", warning: bool = False):
        nonlocal all_passed
        status = "PASS" if passed else ("WARN" if warning else "FAIL")
        if not passed and not warning:
            all_passed = False
        steps.append({"name": name, "status": status, "detail": detail})

    # ── Step 1: Folder exists ─────────────────────────────────────────────
    _step("Submission folder exists", root.exists(),
          str(root) if root.exists() else f"Not found: {root}")

    if not root.exists():
        return {"passed": False, "steps": steps, "submission_root": str(root)}

    # ── Step 2: Manifest valid ────────────────────────────────────────────
    manifest_path = root / "submission_manifest.json"
    if manifest_path.exists():
        try:
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            _step("Manifest JSON valid",  True,
                  f"Product: {manifest.get('full_product_name', '?')}")
        except Exception as e:
            _step("Manifest JSON valid", False, str(e))
    else:
        _step("Manifest JSON valid", False, "submission_manifest.json not found")

    # ── Step 3: Module folders ────────────────────────────────────────────
    for mod in ["Module1", "Module2", "Module3", "Module4", "Module5"]:
        p = root / mod
        _step(f"{mod} folder present", p.exists(),
              f"{len(list(p.rglob('*')))} items" if p.exists() else "Missing")

    # ── Step 4: Key DOCX files non-empty ─────────────────────────────────
    key_docs = [
        ("SmPC",            "Module1/1.3_Product_Information/1.3.1_SmPC/smpc.docx"),
        ("PIL",             "Module1/1.3_Product_Information/1.3.2_PIL/pil.docx"),
        ("Cover Letter",    "Module1/1.1_Cover_Letter/1.1_cover_letter.docx"),
        ("QOS",             "Module2/2.3_QOS/2.3_quality_overall_summary.docx"),
        ("DS General Info", "Module3/3.2_Body_of_Data/3.2.S_Drug_Substance/3.2.S.1_General_Information/3.2.S.1_general_information.docx"),
        ("DS Spec",         "Module3/3.2_Body_of_Data/3.2.S_Drug_Substance/3.2.S.4_Control/3.2.S.4.1_specifications.docx"),
        ("DP Description",  "Module3/3.2_Body_of_Data/3.2.P_Drug_Product/3.2.P.1_Description/3.2.P.1_description_composition.docx"),
        ("DP Spec",         "Module3/3.2_Body_of_Data/3.2.P_Drug_Product/3.2.P.5_Control_of_Drug_Product"),
    ]
    for label, rel_path in key_docs:
        p = root / rel_path
        if p.exists():
            size = p.stat().st_size
            _step(f"{label} DOCX present & non-empty",
                  size > 1000,
                  f"{size:,} bytes",
                  warning=(size <= 1000))
        else:
            _step(f"{label} DOCX present & non-empty", False, "File not found")

    # ── Step 5: Resolved spec exists ─────────────────────────────────────
    import re as _re
    drug_slug = _re.sub(r"[^\w]", "_", drug_name.lower()).strip("_")
    spec_paths = [
        Path("pharmacopoeia_db") / "resolved" / f"{drug_slug}.json",
        Path("pharmacopoeia_db") / "resolved" / f"{drug_slug}_ai_filled.json",
    ]
    spec_found = any(p.exists() for p in spec_paths)
    _step("Resolved spec JSON exists", spec_found,
          str(next((p for p in spec_paths if p.exists()), spec_paths[0])))

    if spec_found:
        spec_path = next(p for p in spec_paths if p.exists())
        try:
            spec = json.loads(spec_path.read_text(encoding="utf-8"))
            fields     = spec.get("fields", {})
            filled     = sum(1 for f in fields.values() if f.get("value"))
            total_f    = len(fields)
            _step("Spec fields populated",
                  filled > (total_f * 0.5),
                  f"{filled}/{total_f} fields resolved",
                  warning=(filled <= total_f * 0.5))
        except Exception as e:
            _step("Spec fields populated", False, str(e))

    # ── Step 6: NAFDAC checklist score ────────────────────────────────────
    validation = validate_submission(root)
    score      = validation["summary"]["score"]
    ready      = validation["summary"]["ready_to_submit"]
    _step("NAFDAC checklist score",
          score >= 70,
          f"{score}% ({validation['summary']['passed']}/{validation['summary']['total']} items)",
          warning=(score < 70))

    # ── Step 7: OpenFDA cache ─────────────────────────────────────────────
    openfda_cache = Path("scrapers_cache") / "openfda" / f"{drug_slug}.json"
    if openfda_cache.exists():
        try:
            od = json.loads(openfda_cache.read_text(encoding="utf-8"))
            _step("OpenFDA data cached", od.get("query_status") == "found",
                  f"Status: {od.get('query_status')}")
        except Exception:
            _step("OpenFDA data cached", False, "Cache file corrupt")
    else:
        _step("OpenFDA data cached", False,
              f"Run: python cli.py drugbank-lookup {drug_name}",
              warning=True)

    # ── Step 8: Pharmacopoeia DB has drug ────────────────────────────────
    bp_dir = Path("pharmacopoeia_db") / "BP" / "json"
    if bp_dir.exists():
        matches = list(bp_dir.glob(f"*{drug_slug[:6]}*.json"))
        _step("Drug in pharmacopoeia DB",
              len(matches) > 0,
              f"{len(matches)} BP match(es) found",
              warning=(len(matches) == 0))
    else:
        _step("Drug in pharmacopoeia DB", False,
              "BP database not built — run: python cli.py build-pharmacopoeia-db",
              warning=True)

    passed_count = sum(1 for s in steps if s["status"] == "PASS")
    warn_count   = sum(1 for s in steps if s["status"] == "WARN")
    fail_count   = sum(1 for s in steps if s["status"] == "FAIL")

    return {
        "passed":           all_passed,
        "submission_root":  str(root),
        "drug_name":        drug_name,
        "steps":            steps,
        "step_summary": {
            "total":  len(steps),
            "passed": passed_count,
            "warn":   warn_count,
            "failed": fail_count,
        },
        "nafdac_score":     score,
        "ready_to_submit":  ready,
        "run_at":           datetime.now().strftime("%Y-%m-%d %H:%M"),
    }