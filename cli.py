"""
NAFDAC Dossier Generator — CLI Entry Point
"""
from nafdac_validator import (
    validate_submission, save_validation_report,
    run_integration_test, NAFDAC_CHECKLIST,
)
from web_scrapers import (
    nafdac_search, nafdac_verify_registration_number,
    who_prequal_check, drugbank_lookup, openfda_lookup,
    fetch_all_external, scraper_cache_stats, clear_scraper_cache,
)
from ai_narrative import (
    tweak_narrative, tweak_docx_section,
    compute_diff, diff_to_terminal, diff_stats, save_diff_docx,
    fill_spec_gaps, apply_gap_fills_to_spec,
    TWEAK_MODES,
)
from pdf_exporter import (
    convert_docx_to_pdf,
    convert_submission_to_pdf,
    find_soffice,
    libreoffice_version,
    PDFExportError,
    BatchResult,
)
from nafdac_structure import (
    SubmissionConfig, build_submission_folder,
    load_manifest, list_submissions, config_from_manifest,
)
from structure_fetcher import (
    fetch_structure, fetch_all_structures,
    get_structure_path, structure_exists,
)
from pubchem_api import (
    query_pubchem, map_to_dossier_fields,
    enrich_bp_database, cache_stats, clear_cache
)

from spec_resolver import (
    resolve_spec, get_flat_spec,
    get_provenance_report, save_resolved_spec,
)
import re
import json
import typer
from module_generators import generate_all_modules
from smpc_pil_generator import generate_smpc, generate_pil
from rich.console import Console
from rich.panel import Panel
from rich.table import Table
from rich.text import Text
from rich import print as rprint
from pathlib import Path
from typing import Optional
from pharmacopoeia_db_builder import build_db, add_single, lookup, db_stats, rebuild_index

app = typer.Typer(
    name="nafdac",
    help="NAFDAC Nigeria Pharmaceutical Dossier Generator",
    add_completion=False,
    rich_markup_mode="rich",
)
console = Console()

def _banner():
    text = Text()
    text.append("NAFDAC", style="bold green")
    text.append(" Dossier Generator", style="bold white")
    text.append(" v0.1.0", style="dim")
    console.print(Panel(text, expand=False))

@app.command()
def ingest(
    input: Path = typer.Option(..., "--input", "-i"),
    drug: str = typer.Option(..., "--drug", "-d"),
    output: Path = typer.Option(Path("templates/"), "--output", "-o"),
):
    """[Phase 2] Ingest an existing dossier and extract reusable templates."""
    _banner()
    suffix = input.suffix.lower() if input.is_file() else ""
    try:
        if input.is_dir() or suffix == ".docx":
            from ingestion.docx_parser import parse_dossier_docx
            parsed = parse_dossier_docx(input, drug_name_hint=drug)
        elif suffix == ".pdf":
            from ingestion.pdf_parser import parse_dossier_pdf
            parsed = parse_dossier_pdf(input, drug_name_hint=drug)
        else:
            console.print("[red]Unsupported file type.[/red]")
            raise typer.Exit(code=1)
    except FileNotFoundError as e:
        console.print(f"[red]{e}[/red]")
        raise typer.Exit(code=1)
    parsed.print_summary()
    from ingestion.template_extractor import extract_templates
    library = extract_templates(parsed, drug_name_hint=drug, templates_root=output)
    library.print_summary()

@app.command(name="new-drug")
def new_drug(
    name: str = typer.Option(..., "--name", "-n"),
    output: Path = typer.Option(Path("config/drug_profile.yaml"), "--output", "-o"),
):
    """[Phase 1] Create a new drug profile YAML."""
    _banner()
    from config_manager import create_blank_drug_profile
    create_blank_drug_profile(name=name, output_path=output)
    console.print(f"[green]Drug profile created at:[/green] {output}")

@app.command()
def generate(
    drug_profile: Path = typer.Option(..., "--drug-profile", "-p"),
    output: Path = typer.Option(Path("output/"), "--output", "-o"),
    skip_ai: bool = typer.Option(False, "--skip-ai"),
):
    """[Phase 4] Generate a complete NAFDAC dossier from a drug profile."""
    _banner()
    console.print("[dim]Phase 4 not yet implemented.[/dim]")

@app.command(name="export-pdf")
def export_pdf(
    path: Path = typer.Option(..., "--path", "-p"),
    keep_docx: bool = typer.Option(True, "--keep-docx/--no-keep-docx"),
):
    """[Phase 4] Convert .docx files to PDF via LibreOffice."""
    _banner()
    console.print("[dim]Phase 4.4 not yet implemented.[/dim]")

@app.command()
def diff(
    original: Path = typer.Option(..., "--original"),
    generated: Path = typer.Option(..., "--generated"),
    export: Optional[Path] = typer.Option(None, "--export"),
):
    """[Phase 6] Show diff between original template and AI-generated output."""
    _banner()
    console.print("[dim]Phase 6.2 not yet implemented.[/dim]")

@app.command(name="add-pharmacopoeia")
def add_pharmacopoeia(
    source: str = typer.Option(..., "--source", "-s"),
    type: str = typer.Option(..., "--type", "-t"),
):
    """[Phase 3] Parse a single BP HTML or USP PDF and add it to the database."""
    _banner()
    if type.upper() not in ("BP", "USP"):
        rprint("[red]ERROR: --type must be BP or USP[/red]")
        raise typer.Exit(1)
    result = add_single(source, type.upper(), verbose=True)
    if "error" in result:
        rprint(f"[red]Failed: {result['error']}[/red]")
        raise typer.Exit(1)
    missing = [w for w in result.get("parse_warnings", []) if w.startswith("MISSING")]
    if missing:
        rprint("[yellow]Missing fields (PubChem will fill):[/yellow]")
        for w in missing:
            rprint(f"  [dim]{w}[/dim]")
    else:
        rprint("[green]All fields populated[/green]")
    rprint(f"Drug: {result.get('drug_name')}  |  Formula: {result.get('molecular_formula')}")

@app.command(name="build-pharmacopoeia-db")
def build_pharmacopoeia_db(
    type: str = typer.Option("BP", "--type", "-t"),
    resume: bool = typer.Option(True, "--resume/--no-resume"),
    max_files: Optional[int] = typer.Option(None, "--max-files"),
):
    """[Phase 3] Parse all BP/USP monographs and build the local JSON database."""
    _banner()
    if type.upper() not in {"BP", "USP", "ALL"}:
        rprint("[red]ERROR: --type must be BP, USP, or ALL[/red]")
        raise typer.Exit(1)
    rprint(f"[bold cyan]Building {type.upper()} database...[/bold cyan]")
    summary = build_db(source=type.upper(), resume=resume, max_files=max_files, show_progress=True)
    rprint(f"[green]Done[/green] — Processed: {summary['processed']}  Skipped: {summary['skipped']}  Errors: {summary['errors']}")

@app.command(name="lookup-pharmacopoeia")
def lookup_pharmacopoeia(
    drug: str = typer.Argument(...),
    show_fields: bool = typer.Option(False, "--fields", "-f"),
):
    """[Phase 3] Look up a drug in the local pharmacopoeia database."""
    _banner()
    result = lookup(drug)
    if not result["matches"]:
        rprint(f"[red]No match for '{drug}'. Run build-pharmacopoeia-db first.[/red]")
        raise typer.Exit(1)
    rprint(f"[green]Matches:[/green] {result['matches'][:5]}")
    for src, data in [("BP", result["bp_data"]), ("USP", result["usp_data"])]:
        if not data:
            continue
        rprint(f"\n[bold]{src}[/bold]")
        table = Table(show_header=False, box=None)
        table.add_column("Field", style="cyan")
        table.add_column("Value")
        fields = ["drug_name","edition","molecular_formula","description","storage"]
        if show_fields:
            fields += ["assay","related_substances","dissolution","loss_on_drying","uniformity"]
        for f in fields:
            v = data.get(f)
            table.add_row(f, str(v)[:80] if v else "[dim]—[/dim]")
        console.print(table)

@app.command(name="pharmacopoeia-stats")
def pharmacopoeia_stats():
    """[Phase 3] Show statistics about the local pharmacopoeia database."""
    _banner()
    s = db_stats()
    table = Table(show_header=False, box=None)
    table.add_column("Field", style="cyan")
    table.add_column("Value")
    table.add_row("Total drugs", str(s["total_drugs"]))
    table.add_row("BP monographs", str(s["bp_count"]))
    table.add_row("USP monographs", str(s["usp_count"]))
    table.add_row("Both sources", str(s["both_sources"]))
    table.add_row("Last updated", str(s["last_updated"]))
    console.print(table)
    if s["total_drugs"] == 0:
        rprint("[yellow]Empty. Run: python cli.py build-pharmacopoeia-db --type BP[/yellow]")
 
# ─────────────────────────────────────────────
# COMMAND: pubchem-lookup
# ─────────────────────────────────────────────
@app.command(name="pubchem-lookup")
def pubchem_lookup(
    drug: str = typer.Argument(..., help="Drug name to look up on PubChem"),
    no_cache: bool = typer.Option(False, "--no-cache", help="Force fresh API query"),
    fields: bool = typer.Option(False, "--fields", "-f", help="Show all fields"),
):
    """
    [Phase 3] Query PubChem for chemical property data on a drug.

    Examples:
      python cli.py pubchem-lookup metformin
      python cli.py pubchem-lookup "metformin hydrochloride" --fields
      python cli.py pubchem-lookup paracetamol --no-cache
    """
    _banner()
    rprint(f"[bold cyan]PubChem lookup: {drug}[/bold cyan]")
    if no_cache:
        rprint("[yellow]Skipping cache — fresh API query[/yellow]")

    result = query_pubchem(drug, use_cache=not no_cache)
    mapped = map_to_dossier_fields(result)

    if result["query_status"] == "not_found":
        rprint(f"[red]✗ Not found on PubChem: '{drug}'[/red]")
        raise typer.Exit(1)

    cached_label = "[dim](cached)[/dim]" if result.get("cached") else "[green](fresh)[/green]"
    rprint(f"\n[green]✓ Found[/green] {cached_label}  →  CID {result['cid']}")

    table = Table(show_header=False, box=None, padding=(0, 2))
    table.add_column("Field", style="cyan", width=22)
    table.add_column("Value", style="white")

    always_show = [
        ("Molecular formula", mapped.get("molecular_formula")),
        ("Molecular weight",  f"{mapped.get('molecular_weight')} g/mol" if mapped.get("molecular_weight") else None),
        ("IUPAC name",        mapped.get("iupac_name")),
        ("logP",              str(mapped.get("logp")) if mapped.get("logp") is not None else None),
        ("Melting point",     mapped.get("melting_point")),
        ("Solubility",        mapped.get("solubility")),
        ("pKa",               mapped.get("pka")),
    ]

    extra_fields_list = [
        ("InChIKey", mapped.get("inchikey")),
        ("SMILES",   (mapped.get("smiles") or "")[:60] + "…"
                     if mapped.get("smiles") and len(mapped.get("smiles", "")) > 60
                     else mapped.get("smiles")),
        ("Synonyms", ", ".join(result.get("synonyms", [])[:4])),
    ]

    rows = always_show + (extra_fields_list if fields else [])
    for label, value in rows:
        table.add_row(label, str(value) if value else "[dim]—[/dim]")

    console.print(table)
    rprint(f"\n  [dim]PubChem URL: {result['pubchem_url']}[/dim]")
# ─────────────────────────────────────────────
# COMMAND: pubchem-enrich-db
# ─────────────────────────────────────────────
@app.command(name="pubchem-enrich-db")
def pubchem_enrich_db(
    max_drugs: Optional[int] = typer.Option(
        None, "--max-drugs", help="Limit number of drugs (for testing)"),
    no_skip: bool = typer.Option(
        False, "--no-skip", help="Re-enrich even if already has PubChem data"),
):
    """
    [Phase 3] Enrich all BP JSON monographs with PubChem data.

    Adds molecular weight, IUPAC name, SMILES, logP, melting point,
    and solubility to each BP monograph JSON file.

    Examples:
      python cli.py pubchem-enrich-db --max-drugs 10
      python cli.py pubchem-enrich-db
      python cli.py pubchem-enrich-db --no-skip
    """
    _banner()
    rprint("[bold cyan]Enriching BP database with PubChem data...[/bold cyan]")
    if max_drugs:
        rprint(f"[yellow]Test mode: limiting to {max_drugs} drugs[/yellow]")
    rprint("[dim]This may take a while for all 1,825 drugs. "
           "Results are cached so re-runs are fast.[/dim]\n")

    summary = enrich_bp_database(
        max_drugs=max_drugs,
        skip_cached=not no_skip,
    )

    rprint(f"\n[bold green]✓ Enrichment complete[/bold green]")
    rprint(f"  Enriched  : [green]{summary['enriched']}[/green]")
    rprint(f"  Skipped   : [dim]{summary['skipped']}[/dim]  (already had PubChem data)")
    rprint(f"  Not found : [yellow]{summary['not_found']}[/yellow]  (no PubChem entry)")
    rprint(f"  Errors    : [red]{summary['errors']}[/red]")


# ─────────────────────────────────────────────
# COMMAND: pubchem-stats
# ─────────────────────────────────────────────
@app.command(name="pubchem-stats")
def pubchem_stats(
    clear: bool = typer.Option(False, "--clear", help="Clear the entire PubChem cache"),
):
    """
    [Phase 3] Show PubChem cache statistics.

    Example:
      python cli.py pubchem-stats
      python cli.py pubchem-stats --clear
    """
    _banner()

    if clear:
        n = clear_cache()
        rprint(f"[yellow]Cache cleared: {n} files deleted[/yellow]")
        return

    stats = cache_stats()
    rprint("\n[bold cyan]PubChem Cache Statistics[/bold cyan]")

    table = Table(show_header=False, box=None, padding=(0, 2))
    table.add_column("Field", style="cyan")
    table.add_column("Value", style="white")
    table.add_row("Total cached",   str(stats["total"]))
    table.add_row("Found on PubChem", str(stats["found"]))
    table.add_row("Not found",      str(stats["not_found"]))
    table.add_row("Cache directory", stats["cache_dir"])
    console.print(table)

    if stats["total"] == 0:
        rprint("\n[yellow]Cache is empty. Run:[/yellow]")
        rprint("  python cli.py pubchem-lookup metformin")
        rprint("  python cli.py pubchem-enrich-db --max-drugs 10")

# ─────────────────────────────────────────────
# COMMAND: draw-structure  (step 3.3c)
# ─────────────────────────────────────────────
@app.command(name="draw-structure")
def draw_structure(
    drug: str = typer.Argument(..., help="Drug name to fetch structure for"),
    size: int = typer.Option(300, "--size", "-s", help="Image size in pixels (default 300)"),
    overwrite: bool = typer.Option(False, "--overwrite", help="Re-fetch even if image exists"),
    show: bool = typer.Option(False, "--show", help="Open image after saving (Windows)"),
    smiles: Optional[str] = typer.Option(None, "--smiles", help="Use this SMILES instead of PubChem"),
):
    """
    [Phase 3] Fetch or generate a 2D structure image for a drug.

    Tries PubChem PNG first, then falls back to RDKit from SMILES.
    Saves to pharmacopoeia_db/structures/<drug>.png

    Examples:
      python cli.py draw-structure metformin
      python cli.py draw-structure paracetamol --show
      python cli.py draw-structure "metformin hydrochloride" --size 400
      python cli.py draw-structure aspirin --smiles "CC(=O)Oc1ccccc1C(=O)O"
    """
    _banner()
    rprint(f"[bold cyan]Drawing structure: {drug}[/bold cyan]")

    # Check cache for CID
    from pubchem_api import _load_cache
    cached = _load_cache(drug)
    cid = cached.get("cid") if cached else None

    if cid:
        rprint(f"  CID   : {cid}")
    if smiles:
        rprint(f"  SMILES: {smiles[:60]}")
    elif cached and cached.get("smiles"):
        rprint(f"  SMILES: {cached['smiles'][:60]}")

    path = fetch_structure(
        drug,
        cid=cid,
        smiles=smiles or (cached.get("smiles") if cached else None),
        size=size,
        overwrite=overwrite,
    )

    if not path:
        rprint(f"[red]✗ Could not fetch structure for '{drug}'[/red]")
        rprint("[dim]Tips:[/dim]")
        rprint("  1. Run: python cli.py pubchem-lookup \"" + drug + "\" --no-cache")
        rprint("  2. Provide SMILES manually: --smiles \"<smiles_string>\"")
        raise typer.Exit(1)

    rprint(f"[green]✓ Structure saved:[/green] {path}")
    rprint(f"  Size: {path.stat().st_size:,} bytes  |  {size}×{size} px")

    if show:
        import subprocess, sys
        if sys.platform == "win32":
            subprocess.Popen(["start", str(path)], shell=True)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", str(path)])
        else:
            subprocess.Popen(["xdg-open", str(path)])
        rprint("[dim]Opening image...[/dim]")


# ─────────────────────────────────────────────
# COMMAND: fetch-all-structures
# ─────────────────────────────────────────────
@app.command(name="fetch-all-structures")
def fetch_all_structures_cmd(
    max_drugs: Optional[int] = typer.Option(
        None, "--max-drugs", help="Limit number of drugs (for testing)"),
    overwrite: bool = typer.Option(
        False, "--overwrite", help="Re-fetch existing images"),
    size: int = typer.Option(
        300, "--size", "-s", help="Image size in pixels (default 300)"),
):
    """
    [Phase 3] Fetch 2D structure images for all enriched drugs in the database.

    Uses PubChem PNG first, falls back to RDKit from SMILES.
    Saves all images to pharmacopoeia_db/structures/

    Examples:
      python cli.py fetch-all-structures --max-drugs 10
      python cli.py fetch-all-structures
      python cli.py fetch-all-structures --overwrite --size 400
    """
    _banner()
    rprint("[bold cyan]Fetching structure images for all enriched drugs...[/bold cyan]")
    if max_drugs:
        rprint(f"[yellow]Test mode: limiting to {max_drugs} drugs[/yellow]")

    summary = fetch_all_structures(
        max_drugs=max_drugs,
        overwrite=overwrite,
        size=size,
    )

    rprint(f"\n[bold green]✓ Done[/bold green]")
    rprint(f"  Fetched : [green]{summary['fetched']}[/green]")
    rprint(f"  Skipped : [dim]{summary['skipped']}[/dim]")
    rprint(f"  Failed  : [red]{summary['failed']}[/red]")
    rprint(f"\n  Images saved to: [cyan]pharmacopoeia_db/structures/[/cyan]")

    
# ─────────────────────────────────────────────
# COMMAND: resolve-spec  (step 3.4)
# ─────────────────────────────────────────────
@app.command(name="resolve-spec")
def resolve_spec_cmd(
    drug: str = typer.Argument(..., help="Drug name to resolve"),
    coa: Optional[str] = typer.Option(
        None, "--coa", help="Path to COA PDF or DOCX file"),
    no_ai: bool = typer.Option(
        False, "--no-ai", help="Disable AI fallback"),
    save: bool = typer.Option(
        False, "--save", "-s", help="Save resolved spec to JSON"),
    report: bool = typer.Option(
        False, "--report", "-r", help="Print full provenance report"),
    out: Optional[str] = typer.Option(
        None, "--out", help="Output JSON path (requires --save)"),
):
    """
    [Phase 3] Resolve all Module 3 quality fields for a drug.

    Applies the priority chain:
      1st: existing dossier data  2nd: BP  3rd: USP  4th: COA  5th: AI

    Examples:
      python cli.py resolve-spec "metformin hydrochloride"
      python cli.py resolve-spec "metformin hydrochloride" --no-ai
      python cli.py resolve-spec "metformin hydrochloride" --coa coa.pdf --save
      python cli.py resolve-spec paracetamol --report
    """
    _banner()
    rprint(f"[bold cyan]Resolving spec: {drug}[/bold cyan]")
    if coa:
        rprint(f"  COA file: {coa}")
    if no_ai:
        rprint("[yellow]  AI fallback disabled[/yellow]")

    resolved = resolve_spec(
        drug_name=drug,
        coa_path=coa,
        use_ai=not no_ai,
        verbose=True,
    )

    summary = resolved["summary"]

    # ── Summary table ──────────────────────────────────────────────────────────
    rprint(f"\n[bold]Resolution Summary[/bold]")
    table = Table(show_header=True, header_style="bold magenta", box=None)
    table.add_column("Field", style="cyan", width=25)
    table.add_column("Source", style="green", width=12)
    table.add_column("Confidence", width=10)
    table.add_column("Value", style="white")

    def _fmt_val(val):
        if val is None:
            return "[dim]—[/dim]"
        if isinstance(val, dict):
            return str(val)[:60] + "…" if len(str(val)) > 60 else str(val)
        if isinstance(val, list):
            return f"[{len(val)} items]"
        return str(val)[:70]

    source_colours = {
        "dossier":    "green",
        "BP":         "cyan",
        "USP":        "blue",
        "COA":        "yellow",
        "AI":         "magenta",
        "local":      "green",
        "unresolved": "red",
        "input":      "white",
    }

    for field, entry in resolved["fields"].items():
        src   = entry["source"]
        conf  = entry["confidence"]
        color = source_colours.get(src, "white")
        table.add_row(
            field,
            f"[{color}]{src}[/{color}]",
            conf,
            _fmt_val(entry["value"]),
        )

    console.print(table)

    # ── Stats ──────────────────────────────────────────────────────────────────
    rprint(f"\n  Filled  : [green]{summary['filled']}[/green] / {summary['total']}")
    rprint(f"  By source: {summary['by_source']}")

    if summary["empty_fields"]:
        rprint(f"  [red]Empty ({summary['empty']}): {summary['empty_fields']}[/red]")

    # ── Warnings ───────────────────────────────────────────────────────────────
    if resolved["warnings"]:
        rprint(f"\n[yellow]Warnings ({len(resolved['warnings'])}):[/yellow]")
        for w in resolved["warnings"]:
            rprint(f"  • [dim]{w}[/dim]")

    # ── Provenance report ──────────────────────────────────────────────────────
    if report:
        rprint(f"\n[bold]Provenance Report[/bold]")
        rprint(get_provenance_report(resolved))

    # ── Save ───────────────────────────────────────────────────────────────────
    if save:
        path = save_resolved_spec(resolved, output_path=out)
        rprint(f"\n[green]✓ Saved:[/green] {path}")

# ─────────────────────────────────────────────
# COMMAND: new-submission  (step 4.1)
# ─────────────────────────────────────────────
@app.command(name="new-submission")
def new_submission(
    drug: str = typer.Option(..., "--drug", "-d",
        help="Drug name e.g. 'Amlodipine Besylate'"),
    strength: str = typer.Option(..., "--strength", "-s",
        help="Strength e.g. '10mg'"),
    form: str = typer.Option(..., "--form", "-f",
        help="Dosage form e.g. 'Tablets'"),
    applicant: str = typer.Option("APPLICANT NAME", "--applicant",
        help="Applicant company name"),
    manufacturer: str = typer.Option("MANUFACTURER NAME", "--manufacturer",
        help="Manufacturer name"),
    country: str = typer.Option("Nigeria", "--country",
        help="Country of manufacture"),
    brand: str = typer.Option("", "--brand",
        help="Brand name (optional)"),
    inn: str = typer.Option("", "--inn",
        help="INN / generic name (defaults to drug name)"),
    submission_type: str = typer.Option("new_product", "--type",
        help="new_product | variation | renewal"),
    output: Optional[str] = typer.Option(None, "--output", "-o",
        help="Custom output path (default: submissions/<drug_slug>/)"),
    overwrite: bool = typer.Option(False, "--overwrite",
        help="Recreate folder even if it already exists"),
):
    """
    [Phase 4] Create the NAFDAC CTD folder structure for a new submission.

    Creates all Module 1-5 folders and a submission_manifest.json.

    Examples:
      python cli.py new-submission --drug "Amlodipine Besylate" --strength 10mg --form Tablets
      python cli.py new-submission --drug "Metformin Hydrochloride" --strength 500mg --form Tablets --applicant "ABC Pharma Ltd"
    """
    _banner()
    rprint(f"[bold cyan]Creating submission: {drug} {strength} {form}[/bold cyan]")

    config = SubmissionConfig(
        drug_name       = drug,
        strength        = strength,
        dosage_form     = form,
        applicant       = applicant,
        manufacturer    = manufacturer,
        country_of_mfr  = country,
        brand_name      = brand,
        inn             = inn or drug,
        submission_type = submission_type,
        output_root     = output or "",
    )

    rprint(f"  Slug     : [dim]{config.slug}[/dim]")
    rprint(f"  Root     : [cyan]{config.submission_root}[/cyan]")

    result = build_submission_folder(config, overwrite=overwrite)
    manifest = result["manifest"]

    rprint(f"\n[green]✓ Submission folder created[/green]")

    table = Table(show_header=False, box=None, padding=(0, 2))
    table.add_column("Field", style="cyan")
    table.add_column("Value", style="white")
    table.add_row("Product",      manifest["full_product_name"])
    table.add_row("Applicant",    manifest["applicant"])
    table.add_row("Manufacturer", manifest["manufacturer"])
    table.add_row("Type",         manifest["submission_type"])
    table.add_row("Folders",      str(manifest["folders_created"]))
    table.add_row("DOCX pending", str(manifest["docx_pending"]))
    table.add_row("Root",         manifest["root"])
    console.print(table)

    rprint(f"\n[dim]Next step: generate Module documents with:[/dim]")
    rprint(f"  python cli.py generate-modules --submission {manifest['root']}")


# ─────────────────────────────────────────────
# COMMAND: list-submissions
# ─────────────────────────────────────────────
@app.command(name="list-submissions")
def list_submissions_cmd():
    """
    [Phase 4] List all submissions in the submissions folder.

    Example:
      python cli.py list-submissions
    """
    _banner()
    subs = list_submissions()

    if not subs:
        rprint("[yellow]No submissions found. Run:[/yellow]")
        rprint('  python cli.py new-submission --drug "Amlodipine Besylate" --strength 10mg --form Tablets')
        return

    rprint(f"\n[bold cyan]Submissions ({len(subs)})[/bold cyan]\n")
    table = Table(show_header=True, header_style="bold magenta", box=None)
    table.add_column("Product",         style="cyan", width=40)
    table.add_column("Applicant",       width=20)
    table.add_column("Date",            width=12)
    table.add_column("DOCX pending",    width=14)
    table.add_column("Root",            style="dim")

    for s in subs:
        table.add_row(
            s["full_product_name"],
            s["applicant"],
            s["submission_date"],
            str(s["docx_pending"]),
            s["root"],
        )
    console.print(table)

# ─────────────────────────────────────────────
# COMMAND: generate-modules  (step 4.2)
# ─────────────────────────────────────────────
@app.command(name="generate-modules")
def generate_modules_cmd(
    submission: str = typer.Option(
        ..., "--submission", "-s",
        help="Path to the submission folder (from new-submission command)"),
    drug_spec: Optional[str] = typer.Option(
        None, "--drug-spec",
        help="Drug name to use for loading resolved spec (defaults to drug in manifest)"),
):
    """
    [Phase 4] Generate all Module DOCX documents for a submission.

    Generates: Cover Letter, QOS, 3.2.S.1, 3.2.S.4, 3.2.P.1, 3.2.P.5
    Uses resolved spec from pharmacopoeia_db/resolved/ automatically.

    Examples:
      python cli.py generate-modules --submission submissions/Amlodipine_Besylate_10mg_Tablets
      python cli.py generate-modules --submission submissions/Amlodipine_Besylate_10mg_Tablets --drug-spec "amlodipine besilate"
    """
    _banner()
    rprint(f"[bold cyan]Generating Module documents[/bold cyan]")
    rprint(f"  Submission: [cyan]{submission}[/cyan]")

    from pathlib import Path
    if not Path(submission).exists():
        rprint(f"[red]✗ Submission folder not found: {submission}[/red]")
        rprint("[dim]Run: python cli.py new-submission --drug ... first[/dim]")
        raise typer.Exit(1)

    summary = generate_all_modules(
        submission_root    = submission,
        drug_name_for_spec = drug_spec,
    )

    rprint(f"\n[bold green]✓ Generation complete[/bold green]")
    rprint(f"  Generated : [green]{summary['generated']}[/green]")
    rprint(f"  Skipped   : [dim]{summary['skipped']}[/dim]  (not yet implemented)")
    rprint(f"  Failed    : [red]{summary['failed']}[/red]")

    if summary["files"]:
        rprint(f"\n[bold]Generated files:[/bold]")
        for f in summary["files"]:
            rprint(f"  [cyan]{f}[/cyan]")

# ─────────────────────────────────────────────
# COMMAND: generate-smpc  (step 4.3)
# ─────────────────────────────────────────────
@app.command(name="generate-smpc")
def generate_smpc_cmd(
    submission: str = typer.Option(
        ..., "--submission", "-s",
        help="Path to the submission folder"),
    drug_spec: Optional[str] = typer.Option(
        None, "--drug-spec",
        help="Drug name for loading resolved spec"),
):
    """
    [Phase 4] Generate the SmPC (Summary of Product Characteristics).

    Saves to Module1/1.3_Product_Information/1.3.1_SmPC/smpc.docx

    Examples:
      python cli.py generate-smpc --submission submissions/Amlodipine_Besylate_10mg_Tablets --drug-spec "amlodipine besilate"
      python cli.py generate-smpc --submission submissions/Metformin_Hydrochloride_500mg_Tablets
    """
    _banner()
    rprint(f"[bold cyan]Generating SmPC[/bold cyan]")
    rprint(f"  Submission: [cyan]{submission}[/cyan]")

    from pathlib import Path
    if not Path(submission).exists():
        rprint(f"[red]✗ Submission folder not found: {submission}[/red]")
        raise typer.Exit(1)

    ok = generate_smpc(submission, drug_name_for_spec=drug_spec)

    if ok:
        out = Path(submission) / "Module1" / "1.3_Product_Information" / "1.3.1_SmPC" / "smpc.docx"
        rprint(f"\n[green]✓ SmPC generated:[/green] {out}")
        rprint(f"  Size: {out.stat().st_size:,} bytes")
    else:
        rprint("[red]✗ SmPC generation failed[/red]")
        raise typer.Exit(1)


# ─────────────────────────────────────────────
# COMMAND: generate-pil  (step 4.3)
# ─────────────────────────────────────────────
@app.command(name="generate-pil")
def generate_pil_cmd(
    submission: str = typer.Option(
        ..., "--submission", "-s",
        help="Path to the submission folder"),
    drug_spec: Optional[str] = typer.Option(
        None, "--drug-spec",
        help="Drug name for loading resolved spec"),
):
    """
    [Phase 4] Generate the PIL (Patient Information Leaflet).

    Saves to Module1/1.3_Product_Information/1.3.2_PIL/pil.docx

    Examples:
      python cli.py generate-pil --submission submissions/Amlodipine_Besylate_10mg_Tablets --drug-spec "amlodipine besilate"
      python cli.py generate-pil --submission submissions/Metformin_Hydrochloride_500mg_Tablets
    """
    _banner()
    rprint(f"[bold cyan]Generating PIL[/bold cyan]")
    rprint(f"  Submission: [cyan]{submission}[/cyan]")

    from pathlib import Path
    if not Path(submission).exists():
        rprint(f"[red]✗ Submission folder not found: {submission}[/red]")
        raise typer.Exit(1)

    ok = generate_pil(submission, drug_name_for_spec=drug_spec)

    if ok:
        out = Path(submission) / "Module1" / "1.3_Product_Information" / "1.3.2_PIL" / "pil.docx"
        rprint(f"\n[green]✓ PIL generated:[/green] {out}")
        rprint(f"  Size: {out.stat().st_size:,} bytes")
    else:
        rprint("[red]✗ PIL generation failed[/red]")
        raise typer.Exit(1)


        
# ─────────────────────────────────────────────
# COMMAND: export-pdf  (step 4.4)
# ─────────────────────────────────────────────
@app.command(name="export-pdf")
def export_pdf(
    docx: str = typer.Argument(..., help="Path to a .docx file to convert"),
    output_dir: Optional[str] = typer.Option(
        None, "--output-dir", "-o",
        help="Directory for the PDF (default: same folder as the .docx)"),
    soffice: Optional[str] = typer.Option(
        None, "--soffice",
        help="Path to soffice.exe (auto-detected if omitted)"),
    timeout: int = typer.Option(120, "--timeout", help="Seconds before giving up"),
):
    """[Phase 4] Convert a single DOCX to PDF via LibreOffice headless.

    Example:
      python cli.py export-pdf submissions/Amlodipine_Besylate_10mg_Tablets/Module1/1.3_Product_Information/1.3.1_SmPC/smpc.docx
    """
    from pdf_exporter import convert_docx_to_pdf, PDFExportError
    _banner()
    rprint(f"[bold cyan]Exporting to PDF[/bold cyan]")
    rprint(f"  Source: [cyan]{docx}[/cyan]")
    try:
        pdf_path = convert_docx_to_pdf(
            Path(docx),
            output_dir=Path(output_dir) if output_dir else None,
            soffice=soffice,
            timeout=timeout,
        )
        rprint(f"\n[green]✓ PDF exported:[/green] {pdf_path}")
        rprint(f"  Size: {pdf_path.stat().st_size:,} bytes")
    except FileNotFoundError as exc:
        rprint(f"[red]✗ File not found:[/red] {exc}")
        raise typer.Exit(1)
    except PDFExportError as exc:
        rprint(f"[red]✗ Export failed:[/red] {exc}")
        raise typer.Exit(1)


@app.command(name="export-pdf-batch")
def export_pdf_batch(
    submission: str = typer.Argument(..., help="Path to the submission root folder"),
    mirror: bool = typer.Option(
        True, "--mirror/--no-mirror",
        help="Mirror DOCX tree under a PDF/ subfolder (default: True)"),
    pdf_subdir: str = typer.Option(
        "PDF", "--pdf-subdir",
        help="Name of the mirrored PDF subfolder (default: PDF)"),
    soffice: Optional[str] = typer.Option(
        None, "--soffice", help="Path to soffice.exe (auto-detected if omitted)"),
    timeout: int = typer.Option(120, "--timeout", help="Per-file timeout in seconds"),
    skip_existing: bool = typer.Option(
        True, "--skip-existing/--overwrite",
        help="Skip files whose PDF already exists (default: skip)"),
):
    """[Phase 4] Convert ALL .docx files in a submission folder to PDF.

    PDFs are placed under <submission>/PDF/ mirroring the original tree,
    or next to each source file if --no-mirror is passed.

    Examples:
      python cli.py export-pdf-batch submissions/Amlodipine_Besylate_10mg_Tablets
      python cli.py export-pdf-batch submissions/Amlodipine_Besylate_10mg_Tablets --no-mirror
    """
    from pdf_exporter import (
        convert_submission_to_pdf, find_soffice,
        libreoffice_version, PDFExportError, BatchResult
    )
    _banner()
    rprint(f"[bold cyan]Batch PDF Export[/bold cyan]")
    rprint(f"  Submission: [cyan]{submission}[/cyan]")

    ver = libreoffice_version(soffice)
    if ver:
        rprint(f"  LibreOffice: [dim]{ver}[/dim]")
    else:
        rprint("[red]✗ LibreOffice not detected.[/red]")
        rprint("  Install from https://www.libreoffice.org/download/")
        rprint("  or set: $env:LIBREOFFICE_PATH = \"C:\\path\\to\\soffice.exe\"")
        raise typer.Exit(1)

    if mirror:
        rprint(f"  PDF output: [cyan]{submission}\\{pdf_subdir}\\[/cyan]")
    else:
        rprint("  PDF output: next to each source .docx")

    try:
        result: BatchResult = convert_submission_to_pdf(
            submission_dir=Path(submission),
            mirror_structure=mirror,
            pdf_subdir=pdf_subdir if mirror else None,
            soffice=soffice,
            timeout=timeout,
            skip_existing=skip_existing,
        )
    except (FileNotFoundError, PDFExportError) as exc:
        rprint(f"[red]✗ Batch export failed:[/red] {exc}")
        raise typer.Exit(1)

    rprint(f"\n  Total .docx found : {result.total}")
    rprint(f"  [green]✓ Converted       : {result.success_count}[/green]")

    if result.failure_count:
        rprint(f"  [red]✗ Failed          : {result.failure_count}[/red]")
        for docx_path, err in result.failed:
            rprint(f"    • {docx_path.name}: {err}")
    else:
        rprint("  [green]All files converted successfully.[/green]")

    if result.converted:
        rprint(f"\n[bold]Output files:[/bold]")
        for pdf in result.converted:
            size = pdf.stat().st_size if pdf.exists() else 0
            rprint(f"  [cyan]{pdf}[/cyan]  ({size:,} bytes)")

# ─────────────────────────────────────────────────────────────────────────────
# COMMAND: nafdac-search  (step 5.1a)
# ─────────────────────────────────────────────────────────────────────────────
@app.command(name="nafdac-search")
def nafdac_search_cmd(
    drug: str = typer.Argument(..., help="Drug name to search in NAFDAC Greenbook"),
    category: str = typer.Option("Drugs", "--category", "-c",
        help="Product category: Drugs | Vaccines and Biologics | Medical devices | Veterinary | Herbals and Nutraceuticals | Disinfectants"),
    no_cache: bool = typer.Option(False, "--no-cache", help="Force fresh query"),
    full: bool = typer.Option(False, "--full", "-f", help="Show all listing fields"),
):
    """
    [Phase 5] Search NAFDAC Greenbook for registered products.

    Checks registration status and lists all matching products.

    Examples:
      python cli.py nafdac-search amlodipine
      python cli.py nafdac-search "metformin hydrochloride" --full
      python cli.py nafdac-search paracetamol --no-cache
    """
    _banner()
    rprint(f"[bold cyan]NAFDAC Greenbook Search: {drug}[/bold cyan]")

    result = nafdac_search(drug, category=category, use_cache=not no_cache)

    cached_label = "[dim](cached)[/dim]" if result.get("cached") else "[green](live)[/green]"

    if result["query_status"] == "error":
        rprint(f"[red]✗ Query failed — check your internet connection[/red]")
        raise typer.Exit(1)

    # Registration status banner
    if result["registration_status"] == "registered":
        rprint(f"\n[green]✓ REGISTERED[/green] {cached_label}  — {result['total']} product(s) found")
    else:
        rprint(f"\n[yellow]✗ NOT REGISTERED[/yellow] {cached_label}  — No NAFDAC listings found for '{drug}'")

    if not result["listings"]:
        rprint("\n[dim]Tip: Try a shorter name, or check spelling vs. INN.[/dim]")
        return

    rprint(f"\n[bold]Registered Products[/bold]")
    table = Table(show_header=True, header_style="bold magenta", box=None)
    table.add_column("Product Name",  style="cyan",  width=30)
    table.add_column("NRN",           style="green", width=16)
    table.add_column("Form",          width=14)
    table.add_column("Strength",      width=12)
    table.add_column("Applicant",     width=24)
    if full:
        table.add_column("Approval Date", width=14)
        table.add_column("Status",         width=10)

    for p in result["listings"]:
        row = [
            p.get("product_name", "")[:30],
            p.get("nrn", ""),
            p.get("form", ""),
            p.get("strength", ""),
            p.get("applicant", "")[:24],
        ]
        if full:
            row += [p.get("approval_date", ""), p.get("status", "")]
        table.add_row(*row)

    console.print(table)
    rprint(f"\n[dim]Source: NAFDAC Greenbook — greenbook.nafdac.gov.ng[/dim]")


# ─────────────────────────────────────────────────────────────────────────────
# COMMAND: nafdac-verify  (step 5.1b)
# ─────────────────────────────────────────────────────────────────────────────
@app.command(name="nafdac-verify")
def nafdac_verify_cmd(
    nrn: str = typer.Argument(..., help="NAFDAC Registration Number e.g. A4-1234"),
):
    """
    [Phase 5] Verify a specific NAFDAC Registration Number (NRN).

    Examples:
      python cli.py nafdac-verify A4-1234
      python cli.py nafdac-verify "04-1234"
    """
    _banner()
    rprint(f"[bold cyan]Verifying NRN: {nrn}[/bold cyan]")

    result = nafdac_verify_registration_number(nrn)

    if result.get("valid"):
        rprint(f"\n[green]✓ VALID[/green] — NRN found in NAFDAC Greenbook")
        if result.get("product"):
            p = result["product"]
            table = Table(show_header=False, box=None, padding=(0, 2))
            table.add_column("Field", style="cyan")
            table.add_column("Value")
            for field, label in [
                ("product_name", "Product"),
                ("nrn",          "NRN"),
                ("form",         "Dosage Form"),
                ("strength",     "Strength"),
                ("applicant",    "Applicant"),
                ("approval_date","Approval Date"),
                ("status",       "Status"),
            ]:
                val = p.get(field, "")
                if val:
                    table.add_row(label, str(val))
            console.print(table)
    else:
        rprint(f"\n[red]✗ NOT FOUND[/red] — NRN '{nrn}' not found in NAFDAC Greenbook")
        rprint("[dim]This may mean the product is unregistered, expired, or the NRN format is incorrect.[/dim]")


# ─────────────────────────────────────────────────────────────────────────────
# COMMAND: who-prequal  (step 5.2)
# ─────────────────────────────────────────────────────────────────────────────
@app.command(name="who-prequal")
def who_prequal_cmd(
    drug: str = typer.Argument(..., help="Drug name to check"),
    no_cache: bool = typer.Option(False, "--no-cache", help="Force fresh query"),
    full: bool = typer.Option(False, "--full", "-f", help="Show all listing details"),
):
    """
    [Phase 5] Check WHO Prequalification status for a drug (FPP + API).

    Checks both the Finished Pharmaceutical Product list and the
    Active Pharmaceutical Ingredient list.

    Examples:
      python cli.py who-prequal amlodipine
      python cli.py who-prequal "metformin hydrochloride" --full
      python cli.py who-prequal rifampicin --no-cache
    """
    _banner()
    rprint(f"[bold cyan]WHO Prequalification Check: {drug}[/bold cyan]")

    result = who_prequal_check(drug, use_cache=not no_cache)
    cached_label = "[dim](cached)[/dim]" if result.get("cached") else "[green](live)[/green]"

    rprint(f"\n{cached_label}")
    rprint()

    # FPP status
    if result["fpp_prequalified"]:
        rprint(f"  [green]✓ FPP Prequalified[/green] — {len(result['fpp_listings'])} listing(s)")
    else:
        rprint(f"  [yellow]✗ FPP Not prequalified[/yellow] — not found on WHO FPP list")

    # API status
    if result["api_prequalified"]:
        rprint(f"  [green]✓ API Prequalified[/green] — {len(result['api_listings'])} listing(s)")
    else:
        rprint(f"  [yellow]✗ API Not prequalified[/yellow] — not found on WHO API list")

    if full:
        for label, listings in [("FPP", result["fpp_listings"]), ("API", result["api_listings"])]:
            if listings:
                rprint(f"\n[bold]{label} Listings[/bold]")
                for item in listings[:10]:
                    if isinstance(item, dict):
                        for k, v in item.items():
                            if v and k != "_cached_at":
                                rprint(f"  [cyan]{k}[/cyan]: {str(v)[:80]}")
                    rprint()

    rprint(f"\n[dim]Source: {result.get('who_prequal_url')}[/dim]")


# ─────────────────────────────────────────────────────────────────────────────
# COMMAND: drugbank-lookup  (step 5.3)
# ─────────────────────────────────────────────────────────────────────────────
@app.command(name="drugbank-lookup")
def drugbank_lookup_cmd(
    drug: str = typer.Argument(..., help="Drug name to look up"),
    no_cache: bool = typer.Option(False, "--no-cache", help="Force fresh query"),
    interactions: bool = typer.Option(True, "--interactions/--no-interactions"),
    pk: bool = typer.Option(True, "--pk/--no-pk", help="Show pharmacokinetics"),
):
    """
    [Phase 5] Fetch drug data from OpenFDA: PK, interactions, mechanism of action.

    Note: DrugBank blocks all programmatic access. OpenFDA (api.fda.gov) is used
    instead — it contains the full FDA drug label including clinical pharmacology,
    pharmacokinetics, drug interactions, and mechanism of action.

    Examples:
      python cli.py drugbank-lookup amlodipine
      python cli.py drugbank-lookup "metformin hydrochloride" --no-interactions
    """
    from web_scrapers import openfda_lookup
    _banner()
    rprint(f"[bold cyan]OpenFDA Lookup: {drug}[/bold cyan]")
    rprint(f"[dim]Source: api.fda.gov/drug/label.json[/dim]")

    result = openfda_lookup(drug, use_cache=not no_cache)
    cached_label = "[dim](cached)[/dim]" if result.get("cached") else "[green](live)[/green]"

    if result["query_status"] == "not_found":
        rprint(f"\n[red]✗ Not found on OpenFDA: '{drug}'[/red]")
        rprint("[dim]Try the INN/generic name or a major brand name.[/dim]")
        raise typer.Exit(1)

    if result["query_status"] == "error":
        rprint(f"\n[red]✗ OpenFDA query failed — check your internet connection[/red]")
        raise typer.Exit(1)

    rprint(f"\n[green]✓ Found[/green] {cached_label}")

    # Identity
    table = Table(show_header=False, box=None, padding=(0, 2))
    table.add_column("Field", style="cyan", width=24)
    table.add_column("Value", style="white")
    table.add_row("Generic name",   result.get("generic_name", ""))
    table.add_row("Brand names",    ", ".join(result.get("brand_names", [])))
    table.add_row("Manufacturer",   result.get("manufacturer", ""))
    table.add_row("FDA Appl. No.",  result.get("fda_application_number", "") or "[dim]—[/dim]")
    if result.get("dosage_forms_strengths"):
        table.add_row("Forms/Strengths", result["dosage_forms_strengths"][:80])
    console.print(table)

    # Mechanism of action
    if result.get("mechanism_of_action"):
        rprint(f"\n[bold]Mechanism of Action[/bold]")
        rprint(f"  {result['mechanism_of_action'][:400]}")

    # Indications
    if result.get("indications_and_usage"):
        rprint(f"\n[bold]Indications[/bold]")
        rprint(f"  {result['indications_and_usage'][:300]}")

    # Pharmacokinetics
    if pk and result.get("pharmacokinetics"):
        rprint(f"\n[bold]Pharmacokinetics[/bold]")
        pk_table = Table(show_header=False, box=None, padding=(0, 2))
        pk_table.add_column("Parameter", style="cyan", width=26)
        pk_table.add_column("Value", style="white")
        pk_labels = {
            "bioavailability":        "Bioavailability",
            "half_life":              "Half-life",
            "tmax":                   "Tmax",
            "protein_binding":        "Protein Binding",
            "volume_of_distribution": "Volume of Distribution",
            "clearance":              "Clearance",
        }
        any_pk = False
        for key, label in pk_labels.items():
            val = result["pharmacokinetics"].get(key, "")
            pk_table.add_row(label, str(val) if val else "[dim]—[/dim]")
            if val:
                any_pk = True
        console.print(pk_table)
        if not any_pk:
            rprint("  [dim]PK values not parsed from label — see clinical pharmacology below[/dim]")
        if result.get("clinical_pharmacology"):
            rprint(f"\n[bold]Clinical Pharmacology[/bold]")
            rprint(f"  {result['clinical_pharmacology'][:500]}")

    # Drug interactions
    if interactions and result.get("drug_interactions"):
        rprint(f"\n[bold]Drug Interactions[/bold]")
        rprint(f"  {result['drug_interactions'][:500]}")

    # Warnings
    if result.get("warnings"):
        rprint(f"\n[bold]Warnings[/bold]")
        rprint(f"  {result['warnings'][:300]}")

    rprint(f"\n[dim]Source: OpenFDA — api.fda.gov[/dim]")

# ─────────────────────────────────────────────────────────────────────────────
# COMMAND: fetch-external  (all 3 sources at once)
# ─────────────────────────────────────────────────────────────────────────────
@app.command(name="fetch-external")
def fetch_external_cmd(
    drug: str = typer.Argument(..., help="Drug name"),
    no_cache: bool = typer.Option(False, "--no-cache"),
):
    """
    [Phase 5] Run all external scrapers (NAFDAC + WHO + DrugBank) at once.

    Useful for a quick regulatory intelligence snapshot before generating a dossier.

    Example:
      python cli.py fetch-external amlodipine
      python cli.py fetch-external "metformin hydrochloride" --no-cache
    """
    _banner()
    rprint(f"[bold cyan]External Regulatory Intelligence: {drug}[/bold cyan]")
    rprint("[dim]Querying NAFDAC Greenbook, WHO Prequalification, and DrugBank...[/dim]\n")

    results = fetch_all_external(drug, use_cache=not no_cache)

    # NAFDAC
    nafdac = results.get("nafdac", {})
    if nafdac and not nafdac.get("error"):
        status = "[green]✓ REGISTERED[/green]" if nafdac.get("registration_status") == "registered" else "[yellow]✗ NOT REGISTERED[/yellow]"
        rprint(f"  NAFDAC Greenbook   : {status}  ({nafdac.get('total', 0)} listings)")
    else:
        rprint("  NAFDAC Greenbook   : [red]✗ Query failed[/red]")

    # WHO
    who = results.get("who", {})
    if who and not who.get("error"):
        fpp = "[green]✓[/green]" if who.get("fpp_prequalified") else "[yellow]✗[/yellow]"
        api = "[green]✓[/green]" if who.get("api_prequalified") else "[yellow]✗[/yellow]"
        rprint(f"  WHO Prequalified   : FPP {fpp}   API {api}")
    else:
        rprint("  WHO Prequalification: [red]✗ Query failed[/red]")

  # DrugBank
    db = results.get("drugbank", {})
    if db and db.get("query_status") == "found":
        pk_filled = sum(1 for v in db.get("pharmacokinetics", {}).values() if v)
        rprint(f"  OpenFDA            : [green]✓ Found[/green]  |  PK fields: {pk_filled}/6")
    elif db and db.get("query_status") == "not_found":
        rprint("  OpenFDA            : [yellow]✗ Not found[/yellow]")
    else:
        rprint("  OpenFDA            : [red]✗ Query failed[/red]")

    rprint(f"\n[dim]Run individual commands for full details:[/dim]")
    rprint(f"  python cli.py nafdac-search \"{drug}\"")
    rprint(f"  python cli.py who-prequal \"{drug}\"")
    rprint(f"  python cli.py drugbank-lookup \"{drug}\"")
# ─────────────────────────────────────────────────────────────────────────────
# COMMAND: scraper-stats
# ─────────────────────────────────────────────────────────────────────────────
@app.command(name="scraper-stats")
def scraper_stats_cmd(
    clear: Optional[str] = typer.Option(
        None, "--clear",
        help="Clear cache for: nafdac | who | drugbank | all"),
):
    """
    [Phase 5] Show scraper cache statistics, or clear the cache.

    Examples:
      python cli.py scraper-stats
      python cli.py scraper-stats --clear all
      python cli.py scraper-stats --clear nafdac
    """
    _banner()
    if clear:
        src = None if clear == "all" else clear
        n = clear_scraper_cache(source=src)
        rprint(f"[yellow]Cache cleared: {n} files deleted[/yellow]")
        return

    stats = scraper_cache_stats()
    rprint(f"\n[bold cyan]Scraper Cache Statistics[/bold cyan]")
    table = Table(show_header=False, box=None, padding=(0, 2))
    table.add_column("Source",  style="cyan")
    table.add_column("Cached",  style="white")
    table.add_row("NAFDAC Greenbook",     str(stats["nafdac"]))
    table.add_row("WHO Prequalification", str(stats["who"]))
    table.add_row("DrugBank",             str(stats["drugbank"]))
    table.add_row("Total",                str(stats["total"]))
    table.add_row("Cache dir",            stats["cache_dir"])
    console.print(table)

    # ─────────────────────────────────────────────────────────────────────────────
# COMMAND: tweak-narrative  (step 6.1)
# ─────────────────────────────────────────────────────────────────────────────
@app.command(name="tweak-narrative")
def tweak_narrative_cmd(
    input: str = typer.Option(
        ..., "--input", "-i",
        help="Path to a .docx file OR a plain .txt file with the section text"),
    mode: str = typer.Option(
        "regulatory", "--mode", "-m",
        help="regulatory | expand | tighten | formal"),
    drug: str = typer.Option(
        "", "--drug", "-d", help="Drug name for context"),
    section: str = typer.Option(
        "", "--section", "-s", help="Section name for context e.g. '3.2.S.4'"),
    paragraph: int = typer.Option(
        0, "--paragraph", "-p",
        help="Paragraph index to rewrite (for DOCX input, 0-based, default: 0)"),
    extra: str = typer.Option(
        "", "--extra", "-e", help="Extra instruction appended to the prompt"),
    output: Optional[str] = typer.Option(
        None, "--output", "-o",
        help="Save rewritten text to this path (.txt or .docx)"),
    diff: bool = typer.Option(
        True, "--diff/--no-diff", help="Show terminal diff after rewrite (default: on)"),
    save_diff: Optional[str] = typer.Option(
        None, "--save-diff",
        help="Also save a DOCX diff report to this path"),
):
    """
    [Phase 6] AI-rewrite a dossier section using the Anthropic API.
 
    Modes:
      regulatory — NAFDAC/ICH compliant language
      expand     — add pharmaceutical detail to thin sections
      tighten    — reduce word count while keeping all content
      formal     — translate informal text to scientific tone
 
    Examples:
      python cli.py tweak-narrative --input submissions/Amlodipine_Besylate_10mg_Tablets/Module3/3.2.S.4_Control/3.2.S.4.1_Specifications/specifications.docx --mode regulatory --drug "Amlodipine Besylate" --section "3.2.S.4"
      python cli.py tweak-narrative --input section.txt --mode expand --drug "Metformin HCl" --save-diff diff_report.docx
      python cli.py tweak-narrative --input smpc.docx --mode tighten --paragraph 2
    """
    _banner()
    rprint(f"[bold cyan]Narrative Tweaker[/bold cyan]")
 
    if mode not in TWEAK_MODES:
        rprint(f"[red]✗ Invalid mode '{mode}'. Choose: {list(TWEAK_MODES.keys())}[/red]")
        raise typer.Exit(1)
 
    mode_label = TWEAK_MODES[mode][0]
    rprint(f"  Mode    : [cyan]{mode}[/cyan] ({mode_label})")
    rprint(f"  Input   : {input}")
    if drug:    rprint(f"  Drug    : {drug}")
    if section: rprint(f"  Section : {section}")
 
    input_path = Path(input)
    if not input_path.exists():
        rprint(f"[red]✗ File not found: {input}[/red]")
        raise typer.Exit(1)
 
    # Read input
    if input_path.suffix.lower() == ".docx":
        rprint(f"  Paragraph index: {paragraph}")
        try:
            result = tweak_docx_section(
                input_path, paragraph_index=paragraph,
                mode=mode, drug_name=drug, section_name=section,
                extra_instruction=extra,
            )
        except IndexError as e:
            rprint(f"[red]✗ {e}[/red]")
            raise typer.Exit(1)
    else:
        # Plain text
        text = input_path.read_text(encoding="utf-8").strip()
        result = tweak_narrative(
            text, mode=mode, drug_name=drug, section_name=section,
            extra_instruction=extra,
        )
 
    # Word count change
    wc_before = result["word_count_before"]
    wc_after  = result["word_count_after"]
    wc_delta  = wc_after - wc_before
    sign      = "+" if wc_delta >= 0 else ""
    rprint(f"\n[green]✓ Rewrite complete[/green]  |  "
           f"Words: {wc_before} → {wc_after} ({sign}{wc_delta})")
 
    # Show rewritten text
    rprint(f"\n[bold]Rewritten Text[/bold]")
    rprint(f"[dim]{'─' * 60}[/dim]")
    console.print(result["rewritten"])
    rprint(f"[dim]{'─' * 60}[/dim]")
 
    # Terminal diff
    if diff:
        from ai_narrative import compute_diff, diff_to_terminal, diff_stats
        diffs = compute_diff(result["original"], result["rewritten"])
        stats = diff_stats(diffs)
        rprint(f"\n[bold]Diff[/bold]  [dim](chars added: {stats['added']}  removed: {stats['removed']})[/dim]")
        rprint(f"[dim]{'─' * 60}[/dim]")
        console.print(diff_to_terminal(diffs))
        rprint(f"[dim]{'─' * 60}[/dim]")
 
    # Save rewritten output
    if output:
        out_path = Path(output)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        if out_path.suffix.lower() == ".docx":
            from docx import Document
            doc = Document()
            doc.add_paragraph(result["rewritten"])
            doc.save(out_path)
        else:
            out_path.write_text(result["rewritten"], encoding="utf-8")
        rprint(f"\n[green]✓ Saved rewritten text:[/green] {out_path}")
 
    # Save diff DOCX
    if save_diff:
        diff_path = save_diff_docx(result, Path(save_diff))
        rprint(f"[green]✓ Saved diff report:[/green] {diff_path}")
 
 
# ─────────────────────────────────────────────────────────────────────────────
# COMMAND: narrative-diff  (step 6.2 — standalone diff between two files)
# ─────────────────────────────────────────────────────────────────────────────
@app.command(name="narrative-diff")
def narrative_diff_cmd(
    original: str = typer.Option(
        ..., "--original", help="Path to original .txt or .docx"),
    revised: str = typer.Option(
        ..., "--revised",  help="Path to revised .txt or .docx"),
    paragraph: int = typer.Option(
        0, "--paragraph", "-p",
        help="Paragraph index (for DOCX files, 0-based)"),
    save_diff: Optional[str] = typer.Option(
        None, "--save-diff", help="Save DOCX diff report to this path"),
    drug: str = typer.Option("", "--drug"),
    section: str = typer.Option("", "--section"),
):
    """
    [Phase 6] Show diff between two text files or DOCX paragraphs.
 
    Useful for comparing manually edited vs AI-generated versions.
 
    Examples:
      python cli.py narrative-diff --original original.txt --revised rewritten.txt
      python cli.py narrative-diff --original smpc_v1.docx --revised smpc_v2.docx --paragraph 3 --save-diff diff.docx
    """
    _banner()
    rprint(f"[bold cyan]Narrative Diff[/bold cyan]")
 
    def _read(path_str: str, para_idx: int) -> str:
        p = Path(path_str)
        if not p.exists():
            rprint(f"[red]✗ File not found: {path_str}[/red]")
            raise typer.Exit(1)
        if p.suffix.lower() == ".docx":
            from docx import Document
            doc = Document(p)
            paras = [par for par in doc.paragraphs if par.text.strip()]
            if para_idx >= len(paras):
                rprint(f"[red]✗ Paragraph {para_idx} out of range ({len(paras)} paragraphs)[/red]")
                raise typer.Exit(1)
            return paras[para_idx].text
        return p.read_text(encoding="utf-8").strip()
 
    orig_text = _read(original, paragraph)
    new_text  = _read(revised,  paragraph)
 
    diffs = compute_diff(orig_text, new_text)
    stats = diff_stats(diffs)
 
    rprint(f"\n  Original : {len(orig_text.split())} words")
    rprint(f"  Revised  : {len(new_text.split())} words")
    rprint(f"  Changes  : [green]+{stats['added']} chars[/green]  [red]-{stats['removed']} chars[/red]  {stats['unchanged']} unchanged")
 
    rprint(f"\n[bold]Diff[/bold]")
    rprint(f"[dim]{'─' * 60}[/dim]")
    console.print(diff_to_terminal(diffs))
    rprint(f"[dim]{'─' * 60}[/dim]")
 
    if save_diff:
        fake_result = {
            "original":           orig_text,
            "rewritten":          new_text,
            "mode":               "manual",
            "mode_label":         "manual comparison",
            "drug_name":          drug,
            "section_name":       section,
            "word_count_before":  len(orig_text.split()),
            "word_count_after":   len(new_text.split()),
        }
        diff_path = save_diff_docx(fake_result, Path(save_diff))
        rprint(f"\n[green]✓ Diff report saved:[/green] {diff_path}")
 
 
# ─────────────────────────────────────────────────────────────────────────────
# COMMAND: fill-spec-gaps  (step 6.3)
# ─────────────────────────────────────────────────────────────────────────────
@app.command(name="fill-spec-gaps")
def fill_spec_gaps_cmd(
    drug: str = typer.Argument(..., help="Drug name — must have a resolved spec in pharmacopoeia_db/resolved/"),
    no_openfda: bool = typer.Option(
        False, "--no-openfda", help="Skip OpenFDA cross-reference"),
    save: bool = typer.Option(
        True, "--save/--no-save",
        help="Save updated spec back to pharmacopoeia_db/resolved/ (default: save)"),
    report: bool = typer.Option(
        False, "--report", "-r", help="Show full field-by-field report"),
    dry_run: bool = typer.Option(
        False, "--dry-run",
        help="Show what would be filled without saving"),
):
    """
    [Phase 6] Use AI to fill empty and low-confidence spec fields.
 
    Reads the resolved spec from pharmacopoeia_db/resolved/<drug>.json,
    fills gaps using the Anthropic API (cross-checked against OpenFDA),
    and saves the updated spec back.
 
    Requires ANTHROPIC_API_KEY to be set.
 
    Examples:
      python cli.py fill-spec-gaps "amlodipine besilate"
      python cli.py fill-spec-gaps "metformin hydrochloride" --report
      python cli.py fill-spec-gaps "amlodipine besilate" --dry-run
      python cli.py fill-spec-gaps paracetamol --no-openfda --no-save
    """
    _banner()
    rprint(f"[bold cyan]AI Spec Gap Filler: {drug}[/bold cyan]")
 
    # Load resolved spec
    from spec_resolver import resolve_spec, save_resolved_spec
    drug_slug = re.sub(r"[^\w]", "_", drug.lower()).strip("_")
    spec_path = Path("pharmacopoeia_db") / "resolved" / f"{drug_slug}.json"
 
    if not spec_path.exists():
        rprint(f"[yellow]No resolved spec found at {spec_path}[/yellow]")
        rprint(f"[dim]Running resolve-spec first...[/dim]")
        resolved = resolve_spec(drug_name=drug, use_ai=False, verbose=False)
        fields = resolved.get("fields", {})
    else:
        fields = json.loads(spec_path.read_text(encoding="utf-8")).get("fields", {})
        rprint(f"  Loaded: [cyan]{spec_path}[/cyan]")
 
    # Count what needs filling
    from ai_narrative import FILLABLE_FIELDS, LOW_CONFIDENCE_SOURCES
    empty  = [f for f in FILLABLE_FIELDS if not fields.get(f, {}).get("value")]
    low_c  = [f for f in FILLABLE_FIELDS
               if fields.get(f, {}).get("value")
               and (fields[f].get("source", "") in LOW_CONFIDENCE_SOURCES
                    or fields[f].get("confidence") == "low")]
 
    rprint(f"\n  Empty fields       : [red]{len(empty)}[/red]")
    rprint(f"  Low-confidence     : [yellow]{len(low_c)}[/yellow]")
    rprint(f"  OpenFDA context    : {'disabled' if no_openfda else 'enabled'}")
 
    if not empty and not low_c:
        rprint("\n[green]✓ Spec is fully resolved — nothing to fill.[/green]")
        return
 
    if dry_run:
        rprint(f"\n[yellow]Dry run — would fill:[/yellow]")
        for f in empty:  rprint(f"  [red]EMPTY[/red]    {f}")
        for f in low_c:  rprint(f"  [yellow]LOW-CONF[/yellow] {f}")
        return
 
    rprint(f"\n[dim]Calling Anthropic API ({MODEL})...[/dim]")
 
    try:
        gap_result = fill_spec_gaps(
            drug_name=drug,
            resolved_spec=fields,
            use_openfda=not no_openfda,
        )
    except RuntimeError as e:
        rprint(f"[red]✗ {e}[/red]")
        raise typer.Exit(1)
 
    rprint(f"\n[green]✓ Gap fill complete[/green]")
    rprint(f"  Fields filled      : [green]{gap_result['fields_filled']}[/green]")
    rprint(f"  Empty found        : {gap_result['empty_fields_found']}")
    rprint(f"  Low-conf found     : {gap_result['low_conf_fields_found']}")
 
    if gap_result["cross_check_notes"]:
        rprint(f"\n[bold]Cross-check Notes[/bold]")
        rprint(f"  {gap_result['cross_check_notes']}")
 
    if gap_result["warnings"]:
        rprint(f"\n[yellow]Warnings:[/yellow]")
        for w in gap_result["warnings"]:
            rprint(f"  • {w}")
 
    # Show filled fields
    if report or gap_result["fields_filled"] <= 10:
        rprint(f"\n[bold]Filled Fields[/bold]")
        table = Table(show_header=True, header_style="bold magenta", box=None)
        table.add_column("Field",      style="cyan", width=26)
        table.add_column("Value",      width=35)
        table.add_column("Confidence", width=10)
        table.add_column("Source",     width=12)
        table.add_column("Note",       style="dim")
        for field, fill in gap_result["filled_fields"].items():
            conf  = fill.get("confidence", "")
            conf_col = (
                "[green]high[/green]"   if conf == "high"   else
                "[yellow]medium[/yellow]" if conf == "medium" else
                "[red]low[/red]"
            )
            table.add_row(
                field,
                str(fill.get("value", ""))[:35],
                conf_col,
                fill.get("source", ""),
                fill.get("note", "")[:40],
            )
        console.print(table)
    else:
        rprint(f"  [dim](use --report to see all {gap_result['fields_filled']} filled fields)[/dim]")
 
    # Save
    if save and not dry_run:
        updated_fields = apply_gap_fills_to_spec(fields, gap_result)
        # Wrap back into full resolved spec structure for save
        full_spec = {
            "drug_name": drug,
            "fields":    updated_fields,
            "summary": {
                "filled":      sum(1 for f in updated_fields.values() if f.get("value")),
                "total":       len(updated_fields),
                "ai_gap_fill": gap_result["fields_filled"],
            },
            "warnings": gap_result["warnings"],
        }
        out_path = Path("pharmacopoeia_db") / "resolved" / f"{drug_slug}_ai_filled.json"
        out_path.write_text(
            json.dumps(full_spec, indent=2, ensure_ascii=False), encoding="utf-8"
        )
        rprint(f"\n[green]✓ Saved:[/green] {out_path}")
        rprint("[dim]Original spec unchanged. AI-filled version saved separately.[/dim]")
    
 
# ─────────────────────────────────────────────────────────────────────────────
# COMMAND: validate  (step 7.1 + 7.2)
# ─────────────────────────────────────────────────────────────────────────────
@app.command(name="validate")
def validate_cmd(
    submission: str = typer.Argument(
        ..., help="Path to the submission root folder"),
    save_report: Optional[str] = typer.Option(
        None, "--save-report", "-r",
        help="Save DOCX validation report to this path"),
    show_all: bool = typer.Option(
        False, "--show-all", "-a",
        help="Show all checklist items including N/A (default: show issues only)"),
    strict: bool = typer.Option(
        False, "--strict",
        help="Exit with error code 1 if not ready to submit"),
):
    """
    [Phase 7] Validate a submission against all NAFDAC CTD requirements.
 
    Checks all Module 1-5 folders and documents against the NAFDAC
    registration checklist. Outputs terminal summary + optional DOCX report.
 
    Examples:
      python cli.py validate submissions/Amlodipine_Besylate_10mg_Tablets
      python cli.py validate submissions/Amlodipine_Besylate_10mg_Tablets --save-report validation_report.docx
      python cli.py validate submissions/Amlodipine_Besylate_10mg_Tablets --show-all
    """
    _banner()
    rprint(f"[bold cyan]NAFDAC Dossier Validation[/bold cyan]")
    rprint(f"  Submission: [cyan]{submission}[/cyan]")
 
    sub_path = Path(submission)
    if not sub_path.exists():
        rprint(f"[red]✗ Submission folder not found: {submission}[/red]")
        raise typer.Exit(1)
 
    result = validate_submission(sub_path)
    summ   = result["summary"]
 
    # ── Score banner ───────────────────────────────────────────────────────
    rprint()
    score_color = "green" if summ["score"] >= 80 else "yellow" if summ["score"] >= 50 else "red"
    rprint(f"  [{score_color}]Compliance Score: {summ['score']}%[/{score_color}]  |  "
           f"Validated: {result['validated_at']}")
    rprint()
 
    if summ["ready_to_submit"]:
        rprint("  [bold green]✓ READY TO SUBMIT[/bold green] — All required items present")
    else:
        rprint("  [bold red]✗ NOT READY[/bold red] — See issues below")
 
    # ── Summary counts ─────────────────────────────────────────────────────
    rprint()
    table = Table(show_header=False, box=None, padding=(0, 2))
    table.add_column("Metric", style="cyan", width=28)
    table.add_column("Count",  style="white")
    table.add_row("Total checklist items",   str(summ["total"]))
    table.add_row("[green]✓ Passed[/green]",              str(summ["passed"]))
    table.add_row("[red]✗ Missing (required)[/red]",  str(summ["missing_required"]))
    table.add_row("[yellow]⚠ Empty / stub files[/yellow]",  str(summ["empty"]))
    table.add_row("[dim]– Missing (optional)[/dim]",  str(summ["missing_optional"]))
    table.add_row("[yellow]⚠ Warnings[/yellow]",          str(summ["warnings"]))
    console.print(table)
 
    # ── Missing required ───────────────────────────────────────────────────
    if result["missing_required"]:
        rprint(f"\n[bold red]Missing Required Items ({len(result['missing_required'])})[/bold red]")
        mr_table = Table(show_header=True, header_style="bold red", box=None)
        mr_table.add_column("ID",          style="cyan", width=8)
        mr_table.add_column("Description", width=40)
        mr_table.add_column("Expected Path", style="dim")
        for item in result["missing_required"]:
            mr_table.add_row(item["id"], item["description"], item["path"])
        console.print(mr_table)
 
    # ── Warnings ───────────────────────────────────────────────────────────
    if result["warnings"]:
        rprint(f"\n[bold yellow]Warnings ({len(result['warnings'])})[/bold yellow]")
        for w in result["warnings"]:
            rprint(f"  [yellow]⚠[/yellow] {w}")
 
    # ── Full checklist (optional) ──────────────────────────────────────────
    if show_all:
        rprint(f"\n[bold]Full Checklist[/bold]")
        cl_table = Table(show_header=True, header_style="bold magenta", box=None)
        cl_table.add_column("ID",       style="cyan", width=8)
        cl_table.add_column("Description", width=38)
        cl_table.add_column("Req",      width=4)
        cl_table.add_column("Status",   width=10)
        cl_table.add_column("Note",     style="dim")
 
        status_styles = {
            "PASS":    "[green]PASS[/green]",
            "MISSING": "[red]MISSING[/red]",
            "EMPTY":   "[yellow]EMPTY[/yellow]",
            "WARNING": "[yellow]WARN[/yellow]",
            "N/A":     "[dim]N/A[/dim]",
        }
        for sid, item in result["results"].items():
            cl_table.add_row(
                sid,
                item["description"][:38],
                "✓" if item["required"] else "–",
                status_styles.get(item["status"], item["status"]),
                item.get("note", "")[:40],
            )
        console.print(cl_table)
 
    # ── Save DOCX report ───────────────────────────────────────────────────
    if save_report:
        report_path = save_validation_report(result, Path(save_report))
        rprint(f"\n[green]✓ Validation report saved:[/green] {report_path}")
    else:
        rprint(f"\n[dim]Tip: Save a full DOCX report with --save-report validation.docx[/dim]")
 
    if strict and not summ["ready_to_submit"]:
        raise typer.Exit(1)
 
 
# ─────────────────────────────────────────────────────────────────────────────
# COMMAND: integration-test  (step 7.3)
# ─────────────────────────────────────────────────────────────────────────────
@app.command(name="integration-test")
def integration_test_cmd(
    submission: str = typer.Argument(
        ..., help="Path to submission root folder"),
    drug: str = typer.Option(
        ..., "--drug", "-d",
        help="Drug name (for checking resolved spec and cache)"),
    save_report: Optional[str] = typer.Option(
        None, "--save-report", "-r",
        help="Save DOCX validation report alongside integration results"),
    verbose: bool = typer.Option(
        True, "--verbose/--quiet",
        help="Show per-step details (default: on)"),
):
    """
    [Phase 7] Run the full end-to-end pipeline integration test.
 
    Checks every stage of the pipeline for a submission:
    folder structure, manifests, DOCX files, resolved spec,
    pharmacopoeia DB, OpenFDA cache, and NAFDAC checklist score.
 
    Examples:
      python cli.py integration-test submissions/Amlodipine_Besylate_10mg_Tablets --drug "amlodipine besilate"
      python cli.py integration-test submissions/Metformin_Hydrochloride_500mg_Tablets --drug "metformin hydrochloride" --save-report report.docx
    """
    _banner()
    rprint(f"[bold cyan]End-to-End Integration Test[/bold cyan]")
    rprint(f"  Submission : [cyan]{submission}[/cyan]")
    rprint(f"  Drug       : [cyan]{drug}[/cyan]")
    rprint()
 
    result = run_integration_test(
        submission_root=submission,
        drug_name=drug,
        verbose=verbose,
    )
 
    # ── Step results ───────────────────────────────────────────────────────
    status_icons = {"PASS": "[green]✓[/green]", "WARN": "[yellow]⚠[/yellow]", "FAIL": "[red]✗[/red]"}
    for step in result["steps"]:
        icon   = status_icons.get(step["status"], "?")
        detail = f"  [dim]{step['detail']}[/dim]" if step["detail"] else ""
        rprint(f"  {icon} {step['name']}{detail}")
 
    # ── Overall result ─────────────────────────────────────────────────────
    ss = result["step_summary"]
    rprint()
    rprint(f"  Steps: {ss['total']} total  |  "
           f"[green]{ss['passed']} passed[/green]  |  "
           f"[yellow]{ss['warn']} warnings[/yellow]  |  "
           f"[red]{ss['failed']} failed[/red]")
    rprint(f"  NAFDAC Score: [cyan]{result['nafdac_score']}%[/cyan]")
    rprint()
 
    if result["passed"] and result["ready_to_submit"]:
        rprint("[bold green]✓ ALL CHECKS PASSED — submission is ready[/bold green]")
    elif result["passed"]:
        rprint("[bold yellow]⚠ PIPELINE COMPLETE — review warnings before submitting[/bold yellow]")
    else:
        rprint("[bold red]✗ INTEGRATION TEST FAILED — fix errors above[/bold red]")
        rprint("[dim]Run: python cli.py validate <submission> --show-all for full checklist[/dim]")
 
    # ── Save DOCX validation report if requested ───────────────────────────
    if save_report:
        from nafdac_validator import validate_submission, save_validation_report
        val_result  = validate_submission(submission)
        report_path = save_validation_report(val_result, Path(save_report))
        rprint(f"\n[green]✓ DOCX report saved:[/green] {report_path}")

if __name__ == "__main__":
    app()





