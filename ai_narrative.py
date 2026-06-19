"""
ai_narrative.py — Phase 6
6.1  Narrative tweaker  — AI rewrites dossier sections
6.2  Diff engine        — terminal + DOCX diff output
6.3  Spec gap filler    — AI fills empty/low-confidence spec fields
"""

from __future__ import annotations

import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

import anthropic
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import diff_match_patch as dmp_module

# ── Anthropic client ──────────────────────────────────────────────────────────
def _client() -> anthropic.Anthropic:
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        raise RuntimeError(
            "ANTHROPIC_API_KEY not set.\n"
            "Run: $env:ANTHROPIC_API_KEY = 'sk-ant-...'"
        )
    return anthropic.Anthropic(api_key=key)

MODEL = "claude-sonnet-4-6"

# ═════════════════════════════════════════════════════════════════════════════
# 6.1  NARRATIVE TWEAKER
# ═════════════════════════════════════════════════════════════════════════════

TWEAK_MODES = {
    "regulatory": (
        "regulatory language",
        "Rewrite this dossier section so it fully complies with NAFDAC and ICH CTD guidelines. "
        "Use precise regulatory terminology (e.g. 'complies with', 'not less than', 'not more than', "
        "'as per BP/USP monograph'). Ensure all statements are unambiguous and auditable. "
        "Do not add speculative content — only rewrite what is present."
    ),
    "expand": (
        "section expansion",
        "Expand this dossier section with appropriate pharmaceutical detail. "
        "Add relevant scientific context, standard test methods, acceptance criteria, "
        "and regulatory rationale where they are missing or thin. "
        "All added content must be factually accurate and consistent with ICH Q6A/Q6B guidelines."
    ),
    "tighten": (
        "verbosity reduction",
        "Tighten this dossier section. Remove redundancy, filler phrases, and over-explanation. "
        "Keep all essential regulatory content and data. "
        "Target: reduce word count by 20-40% while preserving full technical meaning."
    ),
    "formal": (
        "formal scientific tone",
        "Rewrite this dossier section in formal scientific and regulatory English. "
        "Replace informal phrasing, colloquialisms, and passive-aggressive hedging with "
        "precise, confident, third-person scientific language appropriate for an NDA/ANDA dossier."
    ),
}


def tweak_narrative(
    text: str,
    mode: str = "regulatory",
    drug_name: str = "",
    section_name: str = "",
    extra_instruction: str = "",
) -> dict:
    """
    AI-rewrite a dossier section.

    Args:
        text:              Original section text.
        mode:              One of: regulatory | expand | tighten | formal
        drug_name:         Drug name for context.
        section_name:      Section heading for context (e.g. '3.2.S.4 Control of Drug Substance').
        extra_instruction: Any additional free-text instruction appended to the prompt.

    Returns:
        {
          "original": str,
          "rewritten": str,
          "mode": str,
          "mode_label": str,
          "drug_name": str,
          "section_name": str,
          "word_count_before": int,
          "word_count_after": int,
          "model": str,
        }
    """
    if mode not in TWEAK_MODES:
        raise ValueError(f"mode must be one of: {list(TWEAK_MODES)}")

    mode_label, instruction = TWEAK_MODES[mode]

    system = (
        "You are a senior regulatory affairs specialist with expertise in pharmaceutical dossier "
        "writing for African and international regulatory submissions (NAFDAC, EMA, FDA, WHO PQ). "
        "You produce ICH CTD-compliant technical writing. "
        "When rewriting, output ONLY the rewritten section text — no preamble, no commentary, "
        "no markdown formatting, no section headers unless they were in the original."
    )

    context_parts = []
    if drug_name:
        context_parts.append(f"Drug: {drug_name}")
    if section_name:
        context_parts.append(f"Section: {section_name}")
    context_line = "  |  ".join(context_parts)

    user_prompt = (
        f"{context_line}\n\n"
        f"Task: {instruction}\n"
        + (f"Additional instruction: {extra_instruction}\n" if extra_instruction else "")
        + f"\n--- ORIGINAL TEXT ---\n{text}\n--- END ---\n\n"
        "Rewritten text:"
    )

    client = _client()
    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=system,
        messages=[{"role": "user", "content": user_prompt}],
    )

    rewritten = response.content[0].text.strip()

    return {
        "original":           text,
        "rewritten":          rewritten,
        "mode":               mode,
        "mode_label":         mode_label,
        "drug_name":          drug_name,
        "section_name":       section_name,
        "word_count_before":  len(text.split()),
        "word_count_after":   len(rewritten.split()),
        "model":              MODEL,
    }


def tweak_docx_section(
    docx_path: Path,
    paragraph_index: int,
    mode: str = "regulatory",
    drug_name: str = "",
    section_name: str = "",
    extra_instruction: str = "",
) -> dict:
    """
    Tweak a specific paragraph (by index) inside a DOCX file.
    Returns the tweak result dict; does NOT save the file (caller decides).
    """
    doc = Document(docx_path)
    paras = [p for p in doc.paragraphs if p.text.strip()]
    if paragraph_index >= len(paras):
        raise IndexError(
            f"Paragraph index {paragraph_index} out of range "
            f"(document has {len(paras)} non-empty paragraphs)"
        )
    original_text = paras[paragraph_index].text
    return tweak_narrative(
        original_text, mode=mode,
        drug_name=drug_name, section_name=section_name,
        extra_instruction=extra_instruction,
    )


# ═════════════════════════════════════════════════════════════════════════════
# 6.2  DIFF ENGINE
# ═════════════════════════════════════════════════════════════════════════════

def compute_diff(original: str, rewritten: str) -> list[tuple[int, str]]:
    """
    Compute character-level diff using diff-match-patch.
    Returns list of (op, text) where op: -1=delete, 0=equal, 1=insert.
    """
    dmp = dmp_module.diff_match_patch()
    diffs = dmp.diff_main(original, rewritten)
    dmp.diff_cleanupSemantic(diffs)
    return diffs


def diff_to_terminal(diffs: list[tuple[int, str]]) -> str:
    """
    Format diffs for Rich terminal display.
    Deletions in red strikethrough, insertions in green.
    """
    parts = []
    for op, text in diffs:
        if op == -1:
            parts.append(f"[red strike]{text}[/red strike]")
        elif op == 1:
            parts.append(f"[green]{text}[/green]")
        else:
            parts.append(text)
    return "".join(parts)


def diff_stats(diffs: list[tuple[int, str]]) -> dict:
    """Count characters added, removed, and unchanged."""
    added = removed = unchanged = 0
    for op, text in diffs:
        n = len(text)
        if op == 1:
            added += n
        elif op == -1:
            removed += n
        else:
            unchanged += n
    return {"added": added, "removed": removed, "unchanged": unchanged}


def save_diff_docx(
    tweak_result: dict,
    output_path: Path,
) -> Path:
    """
    Save a DOCX showing original vs rewritten with tracked-changes-style highlighting.

    Layout:
      Section heading
      [ORIGINAL — highlighted yellow]
      [REWRITTEN — highlighted green]
      Statistics row
    """
    doc = Document()

    # ── Title ─────────────────────────────────────────────────────────────
    title = doc.add_heading("Narrative Diff Report", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    meta = doc.add_paragraph()
    meta.add_run(f"Drug: ").bold = True
    meta.add_run(tweak_result.get("drug_name", "") or "—")
    meta.add_run("    Section: ").bold = True
    meta.add_run(tweak_result.get("section_name", "") or "—")
    meta.add_run("    Mode: ").bold = True
    meta.add_run(tweak_result.get("mode_label", tweak_result.get("mode", "")))
    meta.add_run(f"    Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_paragraph()

    # ── Original ──────────────────────────────────────────────────────────
    orig_heading = doc.add_paragraph()
    orig_run = orig_heading.add_run("ORIGINAL")
    orig_run.bold = True
    orig_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    _set_para_shading(orig_heading, "FFF2CC")  # yellow background

    orig_para = doc.add_paragraph(tweak_result["original"])
    _set_para_shading(orig_para, "FFF9E6")

    doc.add_paragraph()

    # ── Rewritten ─────────────────────────────────────────────────────────
    new_heading = doc.add_paragraph()
    new_run = new_heading.add_run("REWRITTEN")
    new_run.bold = True
    new_run.font.color.rgb = RGBColor(0x37, 0x5A, 0x23)
    _set_para_shading(new_heading, "E2EFDA")  # green background

    new_para = doc.add_paragraph(tweak_result["rewritten"])
    _set_para_shading(new_para, "F0F7EC")

    doc.add_paragraph()

    # ── Inline diff ───────────────────────────────────────────────────────
    diff_heading = doc.add_paragraph()
    diff_heading.add_run("INLINE DIFF  ").bold = True
    diff_heading.add_run("(")
    r_del = diff_heading.add_run("deleted")
    r_del.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    diff_heading.add_run("  /  ")
    r_ins = diff_heading.add_run("inserted")
    r_ins.font.color.rgb = RGBColor(0x37, 0x5A, 0x23)
    diff_heading.add_run(")")

    diffs = compute_diff(tweak_result["original"], tweak_result["rewritten"])
    diff_para = doc.add_paragraph()
    for op, text in diffs:
        run = diff_para.add_run(text)
        if op == -1:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            run.font.strike = True
        elif op == 1:
            run.font.color.rgb = RGBColor(0x37, 0x5A, 0x23)
            run.bold = True

    doc.add_paragraph()

    # ── Stats ─────────────────────────────────────────────────────────────
    stats = diff_stats(diffs)
    wc_before = tweak_result.get("word_count_before", 0)
    wc_after  = tweak_result.get("word_count_after", 0)
    wc_delta  = wc_after - wc_before
    wc_sign   = "+" if wc_delta >= 0 else ""

    stats_para = doc.add_paragraph()
    stats_para.add_run("Statistics:  ").bold = True
    stats_para.add_run(
        f"Words: {wc_before} → {wc_after} ({wc_sign}{wc_delta})    "
        f"Chars added: {stats['added']}    "
        f"Chars removed: {stats['removed']}    "
        f"Unchanged: {stats['unchanged']}"
    )

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def _set_para_shading(para, hex_color: str):
    """Apply background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


# ═════════════════════════════════════════════════════════════════════════════
# 6.3  AI SPEC GAP FILLER
# ═════════════════════════════════════════════════════════════════════════════

# Fields considered "low confidence" when flagged with these source tags
LOW_CONFIDENCE_SOURCES = {"AI", "unresolved", ""}

# Fields the gap filler should attempt to populate
FILLABLE_FIELDS = [
    "description", "solubility", "melting_point", "boiling_point",
    "pka", "logp", "ph", "particle_size", "bulk_density",
    "related_substances", "loss_on_drying", "residue_on_ignition",
    "heavy_metals", "assay", "dissolution", "disintegration",
    "uniformity", "water_content", "microbial_limits",
    "storage", "shelf_life", "packaging",
]


def _build_gap_filler_prompt(
    drug_name: str,
    empty_fields: list[str],
    low_conf_fields: list[dict],
    openfda_context: str,
    existing_spec: dict,
) -> str:
    """Build the prompt for the gap filler."""

    existing_summary = "\n".join(
        f"  {k}: {v['value']}"
        for k, v in existing_spec.items()
        if v.get("value") and v.get("source") not in LOW_CONFIDENCE_SOURCES
    )

    low_conf_summary = "\n".join(
        f"  {f['field']}: {f['current_value']} (source: {f['source']}, confidence: {f['confidence']})"
        for f in low_conf_fields
    )

    empty_list = "\n".join(f"  - {f}" for f in empty_fields)

    return f"""You are a pharmaceutical regulatory expert filling specification gaps for a NAFDAC dossier.

Drug: {drug_name}

ALREADY RESOLVED (from BP/USP/COA — do NOT change these):
{existing_summary or '  (none)'}

LOW CONFIDENCE FIELDS (may need correction):
{low_conf_summary or '  (none)'}

EMPTY FIELDS TO FILL:
{empty_list or '  (none)'}

OPENFDA LABEL DATA (use as cross-reference):
{openfda_context or '  (not available)'}

INSTRUCTIONS:
1. For each empty or low-confidence field, provide the best scientifically accurate value
   based on your knowledge of {drug_name} and standard pharmacopoeial specifications.
2. Cross-check against the OpenFDA data provided above.
3. Flag each value with a confidence level: high | medium | low
4. For values derived from BP/USP monographs, state the source.
5. Do NOT fabricate specific numerical values you are uncertain about — use ranges or
   "refer to BP/USP monograph" instead.
6. Output ONLY valid JSON in this exact format — no preamble, no markdown:

{{
  "filled_fields": {{
    "field_name": {{
      "value": "...",
      "confidence": "high|medium|low",
      "source": "AI|BP|USP|OpenFDA|literature",
      "note": "brief rationale"
    }}
  }},
  "warnings": ["any important caveats"],
  "cross_check_notes": "any discrepancies found between sources"
}}"""


def fill_spec_gaps(
    drug_name: str,
    resolved_spec: dict,
    use_openfda: bool = True,
    openfda_data: Optional[dict] = None,
) -> dict:
    """
    Use AI to fill empty and low-confidence fields in a resolved spec.

    Args:
        drug_name:      Drug name.
        resolved_spec:  Output from spec_resolver.resolve_spec() — the "fields" dict.
        use_openfda:    Whether to fetch OpenFDA data for cross-referencing.
        openfda_data:   Pre-fetched OpenFDA result (skips API call if provided).

    Returns:
        {
          "drug_name": str,
          "filled_fields": { field: { value, confidence, source, note } },
          "empty_fields_found": int,
          "low_conf_fields_found": int,
          "fields_filled": int,
          "warnings": list,
          "cross_check_notes": str,
          "model": str,
        }
    """
    result = {
        "drug_name":             drug_name,
        "filled_fields":         {},
        "empty_fields_found":    0,
        "low_conf_fields_found": 0,
        "fields_filled":         0,
        "warnings":              [],
        "cross_check_notes":     "",
        "model":                 MODEL,
    }

    # ── Classify fields ───────────────────────────────────────────────────
    empty_fields = []
    low_conf_fields = []

    for field in FILLABLE_FIELDS:
        entry = resolved_spec.get(field, {})
        value = entry.get("value")
        source = entry.get("source", "")
        confidence = entry.get("confidence", "")

        if not value:
            empty_fields.append(field)
        elif source in LOW_CONFIDENCE_SOURCES or confidence == "low":
            low_conf_fields.append({
                "field":         field,
                "current_value": str(value)[:100],
                "source":        source,
                "confidence":    confidence,
            })

    result["empty_fields_found"]    = len(empty_fields)
    result["low_conf_fields_found"] = len(low_conf_fields)

    if not empty_fields and not low_conf_fields:
        result["warnings"].append(
            "No empty or low-confidence fields found — spec is fully resolved."
        )
        return result

    # ── OpenFDA context ───────────────────────────────────────────────────
    openfda_context = ""
    if use_openfda:
        if not openfda_data:
            try:
                from web_scrapers import openfda_lookup
                openfda_data = openfda_lookup(drug_name)
            except Exception:
                openfda_data = {}

        if openfda_data and openfda_data.get("query_status") == "found":
            pk = openfda_data.get("pharmacokinetics", {})
            parts = [
                f"Mechanism: {openfda_data.get('mechanism_of_action','')[:200]}",
                f"Clinical pharmacology: {openfda_data.get('clinical_pharmacology','')[:300]}",
                f"PK — half_life: {pk.get('half_life','')}",
                f"PK — protein_binding: {pk.get('protein_binding','')}",
                f"PK — bioavailability: {pk.get('bioavailability','')}",
                f"PK — Vd: {pk.get('volume_of_distribution','')}",
            ]
            openfda_context = "\n".join(p for p in parts if p.split(": ", 1)[-1])

    # ── Call AI ───────────────────────────────────────────────────────────
    prompt = _build_gap_filler_prompt(
        drug_name=drug_name,
        empty_fields=empty_fields,
        low_conf_fields=low_conf_fields,
        openfda_context=openfda_context,
        existing_spec=resolved_spec,
    )

    client = _client()
    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=(
            "You are a pharmaceutical regulatory expert. "
            "Output ONLY valid JSON as instructed. No markdown, no preamble."
        ),
        messages=[{"role": "user", "content": prompt}],
    )

    raw = response.content[0].text.strip()

    # Strip accidental markdown fences
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError as e:
        result["warnings"].append(f"AI returned invalid JSON: {e}. Raw: {raw[:200]}")
        return result

    result["filled_fields"]     = parsed.get("filled_fields", {})
    result["warnings"]         += parsed.get("warnings", [])
    result["cross_check_notes"] = parsed.get("cross_check_notes", "")
    result["fields_filled"]     = len(result["filled_fields"])

    return result


def apply_gap_fills_to_spec(
    resolved_spec: dict,
    gap_fill_result: dict,
    overwrite_low_confidence: bool = True,
) -> dict:
    """
    Merge AI gap fills back into a resolved spec dict.
    Returns updated spec (does not mutate original).
    """
    import copy
    updated = copy.deepcopy(resolved_spec)

    for field, fill in gap_fill_result.get("filled_fields", {}).items():
        existing = updated.get(field, {})
        existing_source = existing.get("source", "")
        existing_conf   = existing.get("confidence", "")

        # Never overwrite high-confidence non-AI fields
        if existing.get("value") and existing_source not in LOW_CONFIDENCE_SOURCES:
            if existing_conf != "low" or not overwrite_low_confidence:
                continue

        updated[field] = {
            "value":      fill.get("value", ""),
            "source":     f"AI ({fill.get('source', 'AI')})",
            "confidence": fill.get("confidence", "low"),
            "note":       fill.get("note", ""),
        }

    return updated