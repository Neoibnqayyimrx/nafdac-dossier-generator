"""
ingestion/docx_parser.py
========================
Extracts structured content from an existing manually-written
NAFDAC dossier in .docx format.

Output is a ParsedDossier dataclass containing:
  - sections: list of DossierSection (heading + paragraphs + tables)
  - metadata: detected title, module guess, total word count
  - raw_tables: all tables as list-of-dicts for easy template extraction

Usage:
    from ingestion.docx_parser import DocxParser
    parser = DocxParser("path/to/dossier.docx")
    result = parser.parse()
    result.print_summary()
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


# ── Data structures ──────────────────────────────────────────────────────────

@dataclass
class ParsedTable:
    """One table extracted from the document."""
    section_heading: str          # Nearest heading above this table
    headers: list[str]            # Column headers (first row if it looks like a header)
    rows: list[dict[str, str]]    # Each row as {header: cell_value}
    raw_rows: list[list[str]]     # Raw rows including header row

    def is_spec_table(self) -> bool:
        """Heuristic: does this look like a specification/test table?"""
        spec_keywords = {"test", "specification", "limit", "method", "result",
                         "assay", "dissolution", "identification", "description"}
        header_text = " ".join(self.headers).lower()
        return bool(spec_keywords & set(header_text.split()))

    def to_jinja_rows(self) -> list[dict[str, str]]:
        """Return rows ready for Jinja2 template rendering."""
        return self.rows


@dataclass
class DossierSection:
    """One logical section of the dossier (heading + its content)."""
    heading: str                          # Section heading text
    heading_level: int                    # 1=H1, 2=H2, 3=H3, 0=no heading
    paragraphs: list[str]                 # All paragraph texts under this heading
    tables: list[ParsedTable]             # Tables that appear under this heading
    module_guess: str = ""                # e.g. "module3_quality"
    raw_text: str = ""                    # Full concatenated text of section

    def word_count(self) -> int:
        return len(self.raw_text.split())

    def is_empty(self) -> bool:
        return not self.raw_text.strip() and not self.tables


@dataclass
class ParsedDossier:
    """Complete parsed representation of one dossier file."""
    source_path: Path
    drug_name_hint: str                   # Passed in by user via CLI --drug flag
    sections: list[DossierSection] = field(default_factory=list)
    raw_tables: list[ParsedTable] = field(default_factory=list)
    detected_title: str = ""
    total_paragraphs: int = 0
    total_tables: int = 0
    total_words: int = 0

    def get_sections_by_module(self, module: str) -> list[DossierSection]:
        return [s for s in self.sections if s.module_guess == module]

    def all_text(self) -> str:
        """Flat dump of all section text — used by template extractor."""
        return "\n\n".join(
            f"### {s.heading}\n{s.raw_text}" for s in self.sections if not s.is_empty()
        )

    def print_summary(self) -> None:
        from rich.console import Console
        from rich.table import Table as RichTable

        console = Console()
        console.print(f"\n[bold green]Parsed:[/bold green] {self.source_path.name}")
        console.print(f"  Title    : {self.detected_title or '(not detected)'}")
        console.print(f"  Drug hint: {self.drug_name_hint}")
        console.print(f"  Sections : {len(self.sections)}")
        console.print(f"  Tables   : {self.total_tables}")
        console.print(f"  Words    : {self.total_words:,}")
        console.print()

        tbl = RichTable(title="Sections", show_lines=True)
        tbl.add_column("Level", style="dim", width=5)
        tbl.add_column("Heading", style="cyan")
        tbl.add_column("Module Guess", style="yellow")
        tbl.add_column("Words", justify="right")
        tbl.add_column("Tables", justify="right")

        for sec in self.sections:
            if not sec.is_empty():
                tbl.add_row(
                    str(sec.heading_level),
                    sec.heading[:60],
                    sec.module_guess or "—",
                    str(sec.word_count()),
                    str(len(sec.tables)),
                )
        console.print(tbl)


# ── Module classifier ────────────────────────────────────────────────────────

# Keywords that hint at which NAFDAC/CTD module a section belongs to.
_MODULE_KEYWORDS: dict[str, list[str]] = {
    "module1_admin": [
        "cover", "application form", "form a", "form b", "form c",
        "administrative", "letter", "declaration", "authorization",
        "power of attorney", "checklist", "fee",
    ],
    "module2_quality": [
        "quality overall summary", "qos", "quality summary",
        "overall summary", "introduction", "drug substance summary",
        "drug product summary",
    ],
    "module3_quality": [
        "drug substance", "drug product", "3.2.s", "3.2.p",
        "general information", "manufacturer", "characterisation",
        "control of drug substance", "reference standard",
        "container closure", "stability", "specification",
        "test procedure", "validation", "batch analysis",
        "excipient", "formulation", "manufacturing process",
        "description of composition",
    ],
    "module4_safety": [
        "nonclinical", "non-clinical", "pharmacology", "toxicology",
        "pharmacokinetics", "genotoxicity", "carcinogenicity",
        "reproductive", "local tolerance", "safety pharmacology",
    ],
    "module5_efficacy": [
        "clinical", "efficacy", "bioequivalence", "bioavailability",
        "clinical study", "clinical trial", "pharmacodynamics",
        "dose response", "main study", "supportive study",
    ],
    "module5_pil_smpc": [
        "summary of product characteristics", "smpc", "spc",
        "patient information leaflet", "pil", "package leaflet",
        "product information", "indications", "contraindications",
        "posology", "method of administration", "side effects",
        "undesirable effects",
    ],
}


def _guess_module(heading: str, paragraph_text: str) -> str:
    combined = (heading + " " + paragraph_text[:200]).lower()
    scores: dict[str, int] = {mod: 0 for mod in _MODULE_KEYWORDS}
    for mod, keywords in _MODULE_KEYWORDS.items():
        for kw in keywords:
            if kw in combined:
                scores[mod] += 1
    best = max(scores, key=lambda m: scores[m])
    return best if scores[best] > 0 else ""


# ── Heading detection ────────────────────────────────────────────────────────

def _heading_level(para: Paragraph) -> int:
    """Return heading level (1-9) or 0 if not a heading."""
    style_name = para.style.name.lower() if para.style else ""
    if style_name.startswith("heading"):
        try:
            return int(style_name.split()[-1])
        except ValueError:
            return 1
    # Fallback: bold + short + Title Case
    text = para.text.strip()
    if (
        text
        and len(text) < 120
        and para.runs
        and all(run.bold for run in para.runs if run.text.strip())
    ):
        return 2
    return 0


# ── Table extractor ──────────────────────────────────────────────────────────

def _extract_table(tbl: Table, section_heading: str) -> ParsedTable:
    """Convert a python-docx Table into a ParsedTable."""
    raw_rows: list[list[str]] = []
    for row in tbl.rows:
        raw_rows.append([cell.text.strip() for cell in row.cells])

    if not raw_rows:
        return ParsedTable(
            section_heading=section_heading,
            headers=[],
            rows=[],
            raw_rows=[],
        )

    # First row is treated as headers if it looks like one
    # (short text, no numeric-only cells, no empty cells)
    first_row = raw_rows[0]
    looks_like_header = (
        all(cell for cell in first_row) and
        not all(re.match(r"^\d+\.?\d*$", cell) for cell in first_row)
    )

    if looks_like_header and len(raw_rows) > 1:
        headers = first_row
        data_rows = raw_rows[1:]
    else:
        headers = [f"col_{i}" for i in range(len(first_row))]
        data_rows = raw_rows

    # Deduplicate merged cells (python-docx repeats merged cell text)
    def dedup_row(row: list[str]) -> list[str]:
        seen_indices: set[int] = set()
        result = []
        for i, cell in enumerate(row):
            if i > 0 and cell == row[i - 1] and i not in seen_indices:
                result.append("")
                seen_indices.add(i)
            else:
                result.append(cell)
        return result

    rows_as_dicts = []
    for raw_row in data_rows:
        deduped = dedup_row(raw_row)
        padded = deduped + [""] * (len(headers) - len(deduped))
        rows_as_dicts.append(dict(zip(headers, padded[: len(headers)])))

    return ParsedTable(
        section_heading=section_heading,
        headers=headers,
        rows=rows_as_dicts,
        raw_rows=raw_rows,
    )


# ── Main parser ──────────────────────────────────────────────────────────────

class DocxParser:
    """
    Parses a NAFDAC dossier .docx file into a structured ParsedDossier.

    Parameters
    ----------
    path : str | Path
        Path to the .docx file.
    drug_name_hint : str
        The drug name/strength as provided via CLI --drug flag.
        Used to guide variable detection in Phase 2.3.
    """

    def __init__(self, path: str | Path, drug_name_hint: str = "") -> None:
        self.path = Path(path)
        self.drug_name_hint = drug_name_hint

        if not self.path.exists():
            raise FileNotFoundError(f"Dossier file not found: {self.path}")
        if self.path.suffix.lower() != ".docx":
            raise ValueError(f"Expected a .docx file, got: {self.path.suffix}")

        self._doc = Document(str(self.path))

    def parse(self) -> ParsedDossier:
        """
        Run the full parse and return a ParsedDossier.
        Iterates document body in order, grouping content under headings.
        """
        result = ParsedDossier(
            source_path=self.path,
            drug_name_hint=self.drug_name_hint,
        )

        # We walk all block-level elements in document order.
        # python-docx exposes paragraphs and tables separately, so we
        # use the underlying XML to preserve ordering.
        body = self._doc.element.body
        current_section = DossierSection(
            heading="[PREAMBLE]",
            heading_level=0,
            paragraphs=[],
            tables=[],
        )
        all_sections: list[DossierSection] = []

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                para = Paragraph(child, self._doc)
                text = para.text.strip()

                if not text:
                    continue

                level = _heading_level(para)

                if level > 0:
                    # Save current section if it has content
                    if current_section.paragraphs or current_section.tables:
                        current_section.raw_text = "\n".join(current_section.paragraphs)
                        current_section.module_guess = _guess_module(
                            current_section.heading, current_section.raw_text
                        )
                        all_sections.append(current_section)

                    current_section = DossierSection(
                        heading=text,
                        heading_level=level,
                        paragraphs=[],
                        tables=[],
                    )

                    # Capture document title from first H1
                    if level == 1 and not result.detected_title:
                        result.detected_title = text

                else:
                    current_section.paragraphs.append(text)

            elif tag == "tbl":
                tbl = Table(child, self._doc)
                parsed_tbl = _extract_table(tbl, current_section.heading)
                current_section.tables.append(parsed_tbl)
                result.raw_tables.append(parsed_tbl)

        # Don't forget the last section
        if current_section.paragraphs or current_section.tables:
            current_section.raw_text = "\n".join(current_section.paragraphs)
            current_section.module_guess = _guess_module(
                current_section.heading, current_section.raw_text
            )
            all_sections.append(current_section)

        result.sections = all_sections
        result.total_paragraphs = sum(len(s.paragraphs) for s in all_sections)
        result.total_tables = len(result.raw_tables)
        result.total_words = sum(s.word_count() for s in all_sections)

        return result


# ── Batch parser (for a folder of docx files) ────────────────────────────────

def parse_folder(folder: str | Path, drug_name_hint: str = "") -> list[ParsedDossier]:
    """
    Parse all .docx files in a folder and return a list of ParsedDossier objects.
    Skips temporary Word files (starting with ~$).
    """
    folder = Path(folder)
    results = []
    docx_files = [f for f in folder.rglob("*.docx") if not f.name.startswith("~$")]

    if not docx_files:
        raise FileNotFoundError(f"No .docx files found in: {folder}")

    for docx_file in sorted(docx_files):
        try:
            parser = DocxParser(docx_file, drug_name_hint=drug_name_hint)
            results.append(parser.parse())
        except Exception as e:
            print(f"[WARN] Skipping {docx_file.name}: {e}")

    return results


# ── CLI-facing function ──────────────────────────────────────────────────────

def parse_dossier_docx(path: str | Path, drug_name_hint: str = "") -> ParsedDossier:
    """
    Entry point called by the CLI ingest command.
    Handles both single files and folders.
    """
    p = Path(path)
    if p.is_dir():
        results = parse_folder(p, drug_name_hint)
        # Merge all into the first result for now;
        # multi-file merging handled properly in Phase 2.5
        primary = results[0]
        for extra in results[1:]:
            primary.sections.extend(extra.sections)
            primary.raw_tables.extend(extra.raw_tables)
            primary.total_tables += extra.total_tables
            primary.total_words += extra.total_words
        return primary
    else:
        parser = DocxParser(p, drug_name_hint)
        return parser.parse()


# ── Self-test ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python docx_parser.py <path_to_dossier.docx> [drug_name]")
        sys.exit(1)

    path = sys.argv[1]
    drug = sys.argv[2] if len(sys.argv) > 2 else ""
    result = parse_dossier_docx(path, drug)
    result.print_summary()
