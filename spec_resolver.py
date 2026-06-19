"""
spec_resolver.py
================
Applies priority logic when populating Module 3 (Quality) sections of the
NAFDAC CTD dossier. For each field, sources are tried in order until a
non-empty value is found.

Priority order
--------------
  1. Existing dossier data  (from ingested CTD DOCX/PDF — step 2)
  2. BP pharmacopoeia JSON  (from pharmacopoeia_db/BP/json/ — step 3.2)
  3. USP pharmacopoeia JSON (from pharmacopoeia_db/USP/json/ — step 3.2)
  4. COA file               (Certificate of Analysis PDF/DOCX — parsed here)
  5. AI                     (Anthropic claude-sonnet — fills remaining gaps)

Output
------
  ResolvedSpec — a dict with two sections:
    "fields"     : {field_name: {"value": ..., "source": ..., "confidence": ...}}
    "summary"    : {total, filled, empty, by_source}

  The "fields" section maps directly to CTD Module 3 sections:
    3.2.S.1  General Information  → drug_name, molecular_formula, molecular_weight,
                                     iupac_name, smiles, description
    3.2.S.3  Characterisation     → melting_point, solubility, logp, pka
    3.2.S.4  Control of Drug Sub  → assay, identification, related_substances,
                                     loss_on_drying, water_content
    3.2.S.6  Container Closure    → storage
    3.2.P.4  Control of Drug Prod → dissolution, uniformity, microbial_limits

Usage
-----
  from spec_resolver import resolve_spec

  spec = resolve_spec(
      drug_name   = "Metformin Hydrochloride",
      dossier_data = {...},   # from ingestion step (optional)
      coa_path     = "path/to/metformin_coa.pdf",  # optional
      use_ai       = True,
  )

  # Access a field
  print(spec["fields"]["assay"]["value"])
  print(spec["fields"]["assay"]["source"])   # "BP"

  # Get all fields filled by AI
  ai_fields = [k for k, v in spec["fields"].items() if v["source"] == "AI"]

CLI (via cli.py)
----------------
  python cli.py resolve-spec --drug "Metformin Hydrochloride"
  python cli.py resolve-spec --drug "Metformin Hydrochloride" --coa metformin_coa.pdf
  python cli.py resolve-spec --drug "Metformin Hydrochloride" --no-ai
"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any, Optional

log = logging.getLogger(__name__)

# ── DB paths ───────────────────────────────────────────────────────────────────
_BP_JSON_DIR      = Path("pharmacopoeia_db") / "BP"  / "json"
_USP_JSON_DIR     = Path("pharmacopoeia_db") / "USP" / "json"
_PUBCHEM_CACHE    = Path("pharmacopoeia_db") / "pubchem_cache"
_STRUCTURES_DIR   = Path("pharmacopoeia_db") / "structures"
_INDEX_FILE       = Path("pharmacopoeia_db") / "index.json"

# ── All Module 3 fields we resolve ────────────────────────────────────────────
_ALL_FIELDS = [
    # 3.2.S.1  General Information
    "drug_name", "molecular_formula", "molecular_weight",
    "iupac_name", "smiles", "inchikey",
    # 3.2.S.3  Characterisation
    "description", "melting_point", "solubility", "logp", "pka",
    # 3.2.S.4  Control of Drug Substance
    "assay", "identification", "related_substances",
    "loss_on_drying", "water_content",
    # 3.2.S.6  Container Closure / Storage
    "storage",
    # 3.2.P  Drug Product
    "dissolution", "uniformity", "microbial_limits",
    # Structure
    "structure_image_path",
]

# Fields where BP/USP store dicts — we serialise to str for AI/COA comparison
_DICT_FIELDS = {"assay", "dissolution", "uniformity", "microbial_limits"}


# ═══════════════════════════════════════════════════════════════════════════════
# FIELD ENTRY HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _entry(value: Any, source: str, confidence: str = "high") -> dict:
    return {"value": value, "source": source, "confidence": confidence}


def _is_empty(value: Any) -> bool:
    """Return True if a value should be treated as missing."""
    if value is None:
        return True
    if isinstance(value, str) and value.strip() in ("", "—", "N/A", "Unknown"):
        return True
    if isinstance(value, list) and len(value) == 0:
        return True
    if isinstance(value, dict):
        # Dict is empty if all values are None
        return all(v is None for v in value.values())
    return False


# ═══════════════════════════════════════════════════════════════════════════════
# SOURCE LOADERS
# ═══════════════════════════════════════════════════════════════════════════════

def _load_index() -> dict:
    if _INDEX_FILE.exists():
        try:
            return json.loads(_INDEX_FILE.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            pass
    return {"entries": {}}


def _find_bp_json(drug_name: str) -> Optional[Path]:
    """Find the BP JSON file for a drug by searching the index and directory."""
    # Try index first
    index = _load_index()
    key   = drug_name.lower().strip()
    entry = index["entries"].get(key)
    if entry and entry.get("bp"):
        p = Path(entry["bp"])
        if p.exists():
            return p

    # Fuzzy: search by normalised filename
    if _BP_JSON_DIR.exists():
        stem = re.sub(r"[^a-z0-9]+", "-", key).strip("-")
        exact = _BP_JSON_DIR / f"{stem}.json"
        if exact.exists():
            return exact

        # Partial match on stem
        for f in _BP_JSON_DIR.glob("*.json"):
            if stem[:12] in f.stem or f.stem[:12] in stem:
                return f
    return None


def _find_usp_json(drug_name: str) -> Optional[Path]:
    """Find the USP JSON file for a drug."""
    index = _load_index()
    key   = drug_name.lower().strip()
    entry = index["entries"].get(key)
    if entry and entry.get("usp"):
        p = Path(entry["usp"])
        if p.exists():
            return p

    if _USP_JSON_DIR.exists():
        stem = re.sub(r"[^a-z0-9]+", "-", key).strip("-")
        exact = _USP_JSON_DIR / f"{stem}.json"
        if exact.exists():
            return exact
    return None


def _load_json_safe(path: Path) -> dict:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as exc:
        log.warning("Failed to load %s: %s", path, exc)
        return {}


def _load_pubchem_cache(drug_name: str) -> dict:
    key  = re.sub(r"[^a-z0-9]+", "_", drug_name.lower().strip()).strip("_")
    path = _PUBCHEM_CACHE / f"{key}.json"
    if path.exists():
        return _load_json_safe(path)
    return {}


# ═══════════════════════════════════════════════════════════════════════════════
# UNIT NORMALISATION
# ═══════════════════════════════════════════════════════════════════════════════

def _normalise_melting_point(value: str) -> str:
    """Convert °F to °C if needed. Returns original string if already °C."""
    if not value:
        return value
    # Pattern: "336 to 342 °F" or "336 °F"
    f_match = re.search(r"(\d+(?:\.\d+)?)\s*(?:to|–|-)\s*(\d+(?:\.\d+)?)\s*°?\s*F", value)
    if f_match:
        lo = (float(f_match.group(1)) - 32) * 5 / 9
        hi = (float(f_match.group(2)) - 32) * 5 / 9
        return f"{lo:.0f}–{hi:.0f} °C (converted from °F)"
    f_single = re.search(r"(\d+(?:\.\d+)?)\s*°?\s*F", value)
    if f_single:
        c = (float(f_single.group(1)) - 32) * 5 / 9
        return f"{c:.0f} °C (converted from °F)"
    return value


def _normalise_solubility(value: str) -> str:
    """Clean up solubility strings. Convert mg/mL NTP values to descriptive."""
    if not value:
        return value
    # "1 to 5 mg/mL at 72 °F (NTP, 1992)" → keep but note temperature
    value = re.sub(r"\(NTP,?\s*\d{4}\)", "(NTP)", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value[:300]


# ═══════════════════════════════════════════════════════════════════════════════
# COA PARSER
# ═══════════════════════════════════════════════════════════════════════════════

def parse_coa(coa_path: str) -> dict:
    """
    Parse a Certificate of Analysis (COA) PDF or DOCX and extract
    quality specification fields.

    Returns a flat dict of field: value pairs extracted from the COA.
    Fields not found are absent from the dict.
    """
    path = Path(coa_path)
    if not path.exists():
        log.warning("COA file not found: %s", coa_path)
        return {}

    suffix = path.suffix.lower()
    text   = ""

    if suffix == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        except Exception as exc:
            log.warning("COA PDF read failed: %s", exc)
            return {}

    elif suffix in (".docx", ".doc"):
        try:
            from docx import Document
            doc  = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs)
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + "\t".join(c.text for c in row.cells)
        except Exception as exc:
            log.warning("COA DOCX read failed: %s", exc)
            return {}
    else:
        log.warning("Unsupported COA format: %s", suffix)
        return {}

    if not text.strip():
        return {}

    result: dict[str, Any] = {}

    # ── Assay ──────────────────────────────────────────────────────────────────
    assay_m = re.search(
        r"[Aa]ssay[:\s]+(\d{2,3}(?:\.\d+)?)\s*%?\s*(?:to|–|-)\s*(\d{2,3}(?:\.\d+)?)\s*%",
        text
    )
    if assay_m:
        result["assay"] = {
            "method": "See COA",
            "limits": {"min": float(assay_m.group(1)), "max": float(assay_m.group(2))},
            "units": "%",
        }
    else:
        # Single value assay result
        single_m = re.search(r"[Aa]ssay[:\s]+(\d{2,3}(?:\.\d+)?)\s*%", text)
        if single_m:
            result["assay"] = {
                "method": "See COA",
                "limits": {"min": None, "max": None},
                "result": float(single_m.group(1)),
                "units": "%",
            }

    # ── Description ───────────────────────────────────────────────────────────
    desc_m = re.search(
        r"(?:[Dd]escription|[Aa]ppearance)[:\s]+([A-Za-z][\w\s,\-\.]+?)(?:\n|$)",
        text
    )
    if desc_m:
        result["description"] = desc_m.group(1).strip()[:200]

    # ── Loss on drying ─────────────────────────────────────────────────────────
    lod_m = re.search(
        r"[Ll]oss\s+on\s+[Dd]rying[:\s]+(?:NMT\s+)?(\d+(?:\.\d+)?)\s*%",
        text
    )
    if lod_m:
        result["loss_on_drying"] = {
            "limit_pct": float(lod_m.group(1)),
            "method": "See COA",
        }

    # ── Water content ──────────────────────────────────────────────────────────
    water_m = re.search(
        r"[Ww]ater\s+(?:[Cc]ontent)?[:\s]+(?:NMT\s+)?(\d+(?:\.\d+)?)\s*%",
        text
    )
    if water_m:
        result["water_content"] = {
            "limit_pct": float(water_m.group(1)),
            "method": "Karl Fischer / See COA",
        }

    # ── Storage ───────────────────────────────────────────────────────────────
    stor_m = re.search(
        r"[Ss]torage[:\s]+([A-Za-z][\w\s,°\-\.]+?)(?:\n|$)",
        text
    )
    if stor_m:
        result["storage"] = stor_m.group(1).strip()[:200]

    # ── Melting point ─────────────────────────────────────────────────────────
    mp_m = re.search(
        r"[Mm]elting\s+[Pp]oint[:\s]+(\d{2,3}(?:[–\-]\d{2,3})?)\s*°?\s*C",
        text
    )
    if mp_m:
        result["melting_point"] = mp_m.group(1) + " °C"

    log.info("COA parsed: found fields %s", list(result.keys()))
    return result


# ═══════════════════════════════════════════════════════════════════════════════
# AI FALLBACK (ANTHROPIC)
# ═══════════════════════════════════════════════════════════════════════════════

def _fill_with_ai(drug_name: str, missing_fields: list[str]) -> dict:
    """
    Use Anthropic Claude to fill remaining missing fields.
    Only called for fields that couldn't be resolved from any other source.

    Returns dict of {field_name: value} for filled fields.
    """
    if not missing_fields:
        return {}

    try:
        import anthropic
    except ImportError:
        log.warning("anthropic SDK not installed. Run: python -m pip install anthropic")
        return {}

    # Build a targeted prompt — only ask for what's missing
    fields_list = "\n".join(f"  - {f}" for f in missing_fields)

    prompt = f"""You are a pharmaceutical data assistant helping prepare a NAFDAC CTD dossier.

Drug: {drug_name}

Please provide the following pharmaceutical quality data fields for this drug.
Return ONLY a JSON object with the field names as keys.
If you are uncertain about a value, set it to null.
Do not include any explanation or preamble — only the JSON object.

Fields needed:
{fields_list}

Field format guidance:
- description: string, e.g. "White or almost white crystalline powder"
- storage: string, e.g. "Store below 25°C, protected from moisture"
- melting_point: string with unit, e.g. "228-232 °C"
- solubility: string, e.g. "Freely soluble in water"
- logp: number, e.g. -2.6
- pka: string, e.g. "pKa = 12.4"
- assay: object with keys: method (string), limits (object with min/max as numbers), units (string)
- loss_on_drying: object with keys: limit_pct (number), method (string)
- water_content: object with keys: limit_pct (number), method (string)
- dissolution: object with keys: apparatus, medium, rpm, time_min, limit_pct
- uniformity: object with keys: method, criterion
- microbial_limits: object with keys: TAMC, TYMC, pathogens (list)

Return only valid JSON."""

    try:
        client = anthropic.Anthropic()
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1000,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = message.content[0].text.strip()

        # Strip markdown code fences if present
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

        ai_data = json.loads(raw)
        log.info("AI filled %d fields for '%s'", len(ai_data), drug_name)
        return ai_data

    except json.JSONDecodeError as exc:
        log.warning("AI response was not valid JSON: %s", exc)
        return {}
    except Exception as exc:
        log.warning("AI call failed: %s", exc)
        return {}


# ═══════════════════════════════════════════════════════════════════════════════
# MASTER RESOLVER
# ═══════════════════════════════════════════════════════════════════════════════

def resolve_spec(
    drug_name: str,
    dossier_data: Optional[dict] = None,
    coa_path: Optional[str] = None,
    use_ai: bool = True,
    verbose: bool = True,
) -> dict:
    """
    Resolve all Module 3 fields for a drug using the priority chain.

    Priority:
      1. dossier_data   (existing CTD dossier — highest trust)
      2. BP JSON        (BP pharmacopoeia database)
      3. USP JSON       (USP pharmacopoeia database)
      4. COA            (Certificate of Analysis)
      5. AI             (Anthropic Claude — lowest trust, fills gaps only)

    Parameters
    ----------
    drug_name    : str   Drug name to resolve.
    dossier_data : dict  Pre-parsed dossier data (from ingestion step 2).
    coa_path     : str   Path to a COA PDF or DOCX file (optional).
    use_ai       : bool  Allow AI fallback for remaining empty fields.
    verbose      : bool  Print resolution log.

    Returns
    -------
    dict with:
      "drug_name"  : str
      "fields"     : {field: {value, source, confidence}}
      "summary"    : {total, filled, empty, empty_fields, by_source}
      "warnings"   : list[str]
    """
    if verbose:
        print(f"\n{'='*60}")
        print(f"  Resolving spec for: {drug_name}")
        print(f"{'='*60}")

    warnings: list[str] = []
    fields: dict[str, dict] = {}

    # ── Load all sources ───────────────────────────────────────────────────────
    dossier = dossier_data or {}

    bp_path = _find_bp_json(drug_name)
    bp_data = _load_json_safe(bp_path) if bp_path else {}
    if bp_path and verbose:
        print(f"  BP  : {bp_path.name}")
    elif verbose:
        print(f"  BP  : not found")

    usp_path = _find_usp_json(drug_name)
    usp_data = _load_json_safe(usp_path) if usp_path else {}
    if usp_path and verbose:
        print(f"  USP : {usp_path.name}")
    elif verbose:
        print(f"  USP : not found")

    coa_data: dict = {}
    if coa_path:
        coa_data = parse_coa(coa_path)
        if verbose:
            print(f"  COA : {len(coa_data)} fields extracted from {coa_path}")

    pubchem = _load_pubchem_cache(drug_name)
    if pubchem and verbose:
        print(f"  PubChem cache: CID {pubchem.get('cid')}")

    # ── Field extraction helpers ───────────────────────────────────────────────

    def _bp(field: str) -> Any:
        """Get a field from BP data. Also checks PubChem cache for chemical fields."""
        val = bp_data.get(field)
        if _is_empty(val) and field in (
            "molecular_formula", "molecular_weight", "iupac_name",
            "smiles", "inchikey", "logp", "melting_point", "solubility", "pka",
        ):
            val = pubchem.get(field)
        return val

    def _usp(field: str) -> Any:
        val = usp_data.get(field)
        if _is_empty(val) and field in (
            "molecular_formula", "molecular_weight", "iupac_name",
            "smiles", "inchikey", "logp", "melting_point", "solubility",
        ):
            val = pubchem.get(field)
        return val

    def _resolve_field(
        field: str,
        dossier_key: Optional[str] = None,
        bp_key: Optional[str] = None,
        usp_key: Optional[str] = None,
        coa_key: Optional[str] = None,
        normalise_fn=None,
    ) -> None:
        """
        Resolve a single field through the priority chain.
        Records the result in `fields` dict.
        """
        dk  = dossier_key or field
        bk  = bp_key      or field
        uk  = usp_key     or field
        ck  = coa_key     or field

        # 1. Dossier
        val = dossier.get(dk)
        if not _is_empty(val):
            if normalise_fn:
                val = normalise_fn(val)
            fields[field] = _entry(val, "dossier", "high")
            return

        # 2. BP
        val = _bp(bk)
        if not _is_empty(val):
            if normalise_fn:
                val = normalise_fn(val)
            fields[field] = _entry(val, "BP", "high")
            return

        # 3. USP
        val = _usp(uk)
        if not _is_empty(val):
            if normalise_fn:
                val = normalise_fn(val)
            fields[field] = _entry(val, "USP", "high")
            return

        # 4. COA
        val = coa_data.get(ck)
        if not _is_empty(val):
            if normalise_fn:
                val = normalise_fn(val)
            fields[field] = _entry(val, "COA", "medium")
            return

        # Mark as unresolved (AI will fill later)
        fields[field] = _entry(None, "unresolved", "none")

    # ── Resolve all fields ─────────────────────────────────────────────────────

    # Drug name — always from input
    fields["drug_name"] = _entry(
        bp_data.get("drug_name") or drug_name,
        "BP" if bp_data.get("drug_name") else "input",
        "high"
    )

    # Chemical identity
    _resolve_field("molecular_formula",
                   normalise_fn=lambda v: str(v).strip() if v else v)
    _resolve_field("molecular_weight",
                   normalise_fn=lambda v: round(float(v), 2) if v else v)
    _resolve_field("iupac_name")
    _resolve_field("smiles")
    _resolve_field("inchikey")

    # Physical characterisation
    _resolve_field("description",
                   bp_key="description", usp_key="description")
    _resolve_field("melting_point",
                   normalise_fn=_normalise_melting_point)
    _resolve_field("solubility",
                   normalise_fn=_normalise_solubility)
    _resolve_field("logp")
    _resolve_field("pka")

    # Quality tests
    _resolve_field("assay")
    _resolve_field("identification")
    _resolve_field("related_substances")
    _resolve_field("loss_on_drying")
    _resolve_field("water_content")
    _resolve_field("storage")

    # Drug product tests
    _resolve_field("dissolution")
    _resolve_field("uniformity")
    _resolve_field("microbial_limits")

    # Structure image
    struct_path = _STRUCTURES_DIR / f"{re.sub(r'[^a-z0-9]+', '-', drug_name.lower()).strip('-')}.png"
    if struct_path.exists():
        fields["structure_image_path"] = _entry(str(struct_path), "local", "high")
    else:
        fields["structure_image_path"] = _entry(None, "unresolved", "none")

    # ── AI fallback ────────────────────────────────────────────────────────────
    unresolved = [f for f, v in fields.items() if v["source"] == "unresolved"]

    if unresolved and use_ai:
        if verbose:
            print(f"\n  AI fallback for {len(unresolved)} unresolved fields: "
                  f"{unresolved}")
        ai_data = _fill_with_ai(drug_name, unresolved)
        for field in unresolved:
            val = ai_data.get(field)
            if not _is_empty(val):
                if field == "melting_point":
                    val = _normalise_melting_point(str(val))
                elif field == "solubility":
                    val = _normalise_solubility(str(val))
                fields[field] = _entry(val, "AI", "low")
                warnings.append(
                    f"Field '{field}' filled by AI — verify before submission."
                )

    # Remaining unresolved after AI
    still_empty = [f for f, v in fields.items() if v["source"] == "unresolved"]
    for f in still_empty:
        warnings.append(f"Field '{f}' could not be resolved from any source.")

    # ── Summary ────────────────────────────────────────────────────────────────
    by_source: dict[str, int] = {}
    for v in fields.values():
        src = v["source"]
        by_source[src] = by_source.get(src, 0) + 1

    filled      = sum(1 for v in fields.values() if v["source"] != "unresolved" and v["value"] is not None)
    empty_count = len(still_empty)
    total       = len(fields)

    summary = {
        "total":        total,
        "filled":       filled,
        "empty":        empty_count,
        "empty_fields": still_empty,
        "by_source":    by_source,
    }

    if verbose:
        print(f"\n  ── Resolution complete ──")
        print(f"     Total fields : {total}")
        print(f"     Filled       : {filled}")
        print(f"     Empty        : {empty_count}")
        print(f"     By source    : {by_source}")
        if warnings:
            print(f"     Warnings ({len(warnings)}):")
            for w in warnings[:5]:
                print(f"       • {w}")

    return {
        "drug_name": drug_name,
        "fields":    fields,
        "summary":   summary,
        "warnings":  warnings,
    }


# ═══════════════════════════════════════════════════════════════════════════════
# FLAT EXPORT (for dossier generators in Phase 4)
# ═══════════════════════════════════════════════════════════════════════════════

def get_flat_spec(resolved: dict) -> dict:
    """
    Return a flat {field: value} dict from a resolved spec.
    Strips provenance metadata — use this for template rendering in Phase 4.
    """
    return {
        field: entry["value"]
        for field, entry in resolved["fields"].items()
    }


def get_provenance_report(resolved: dict) -> str:
    """
    Return a human-readable provenance report for audit purposes.
    Lists each field, its value (truncated), and its source.
    """
    lines = [
        f"Provenance Report: {resolved['drug_name']}",
        "=" * 60,
    ]
    for field, entry in resolved["fields"].items():
        val = entry["value"]
        if val is None:
            val_str = "— (empty)"
        elif isinstance(val, dict):
            val_str = json.dumps(val)[:60] + "…"
        elif isinstance(val, list):
            val_str = f"[{len(val)} items]"
        else:
            val_str = str(val)[:60]
        src  = entry["source"]
        conf = entry["confidence"]
        lines.append(f"  {field:25s} [{src:10s}] ({conf:6s})  {val_str}")

    summary = resolved["summary"]
    lines += [
        "",
        f"Summary: {summary['filled']}/{summary['total']} fields filled",
        f"By source: {summary['by_source']}",
    ]
    return "\n".join(lines)


def save_resolved_spec(resolved: dict, output_path: Optional[str] = None) -> Path:
    """
    Save the resolved spec to a JSON file.
    Default path: pharmacopoeia_db/resolved/<drug_name>.json
    """
    if output_path:
        out = Path(output_path)
    else:
        stem = re.sub(r"[^a-z0-9]+", "-", resolved["drug_name"].lower()).strip("-")
        out  = Path("pharmacopoeia_db") / "resolved" / f"{stem}.json"

    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(
        json.dumps(resolved, indent=2, ensure_ascii=False, default=str),
        encoding="utf-8",
    )
    log.info("Resolved spec saved: %s", out)
    return out