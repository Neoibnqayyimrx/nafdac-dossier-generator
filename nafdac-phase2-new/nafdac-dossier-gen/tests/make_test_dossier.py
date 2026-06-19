"""
Creates a realistic synthetic NAFDAC dossier .docx for testing.
Run: python tests/make_test_dossier.py
"""
from docx import Document
from docx.shared import Pt
from pathlib import Path

def make_test_dossier(output_path: str = "tests/sample_dossier.docx"):
    doc = Document()

    # Module 1 — Admin
    doc.add_heading("MODULE 1: ADMINISTRATIVE INFORMATION", level=1)
    doc.add_heading("1.1 Cover Letter", level=2)
    doc.add_paragraph(
        "This application is submitted by Emzor Pharmaceuticals Limited for the registration "
        "of Amoxicillin 500mg Capsules with the National Agency for Food and Drug Administration "
        "and Control (NAFDAC). The product is manufactured at Plot 20, Murtala Mohammed Airport "
        "Road, Lagos, Nigeria. The manufacturer holds a valid GMP certificate issued on "
        "15 January 2023, valid until 14 January 2025."
    )
    doc.add_heading("1.2 Application Form", level=2)
    doc.add_paragraph(
        "Application Type: New Product Registration\n"
        "Product Name: Amoxicillin 500mg Capsules\n"
        "Dosage Form: Hard Gelatin Capsules\n"
        "Strength: 500mg\n"
        "Applicant: Emzor Pharmaceuticals Limited\n"
        "Date of Submission: 20 March 2023"
    )

    # Module 3 — Quality
    doc.add_heading("MODULE 3: QUALITY", level=1)
    doc.add_heading("3.2.S Drug Substance", level=2)
    doc.add_heading("3.2.S.1 General Information", level=3)
    doc.add_paragraph(
        "INN Name: Amoxicillin\n"
        "Chemical Name: (2S,5R,6R)-6-[(R)-(-)-2-Amino-2-(4-hydroxyphenyl)acetamido]-"
        "3,3-dimethyl-7-oxo-4-thia-1-azabicyclo[3.2.0]heptane-2-carboxylic acid trihydrate\n"
        "Molecular Formula: C16H19N3O5S.3H2O\n"
        "Molecular Weight: 419.45 g/mol\n"
        "CAS Number: 61336-70-7\n"
        "Pharmacopoeia Standard: British Pharmacopoeia (BP) 2023\n"
        "Storage: Store below 25°C, protect from moisture"
    )
    doc.add_heading("3.2.S.4 Control of Drug Substance", level=3)
    doc.add_paragraph(
        "The drug substance Amoxicillin complies with the specifications of the "
        "British Pharmacopoeia 2023. Specification limits are as follows:"
    )

    # Spec table
    tbl = doc.add_table(rows=7, cols=3)
    tbl.style = "Table Grid"
    headers = ["Test", "Specification", "Method"]
    data = [
        ("Description", "White or almost white crystalline powder", "Visual"),
        ("Identification", "Complies with BP identification tests A and B", "BP 2023"),
        ("Assay", "95.0% – 102.0% (dried basis)", "HPLC"),
        ("Water content", "NMT 15.0%", "Karl Fischer"),
        ("Related substances", "Any individual impurity: NMT 2.0%", "HPLC"),
        ("Microbial limit", "Complies with BP limits", "BP 2023"),
    ]
    for i, header in enumerate(headers):
        tbl.rows[0].cells[i].text = header
    for row_idx, (test, spec, method) in enumerate(data, 1):
        tbl.rows[row_idx].cells[0].text = test
        tbl.rows[row_idx].cells[1].text = spec
        tbl.rows[row_idx].cells[2].text = method

    doc.add_heading("3.2.P Drug Product", level=2)
    doc.add_heading("3.2.P.1 Description and Composition", level=3)
    doc.add_paragraph(
        "Amoxicillin 500mg Capsules are hard gelatin capsules filled with white to "
        "off-white powder. Each capsule contains 500mg of Amoxicillin (as trihydrate). "
        "The product has a shelf life of 24 months when stored below 25°C in a dry place. "
        "Pack size: Blister pack of 10 capsules x 10 strips."
    )

    # Composition table
    comp_tbl = doc.add_table(rows=5, cols=3)
    comp_tbl.style = "Table Grid"
    comp_headers = ["Ingredient", "Function", "Amount per Capsule"]
    comp_data = [
        ("Amoxicillin Trihydrate", "Active substance", "574mg (equiv. to 500mg Amoxicillin)"),
        ("Magnesium Stearate", "Lubricant", "5mg"),
        ("Colloidal Silicon Dioxide", "Glidant", "2mg"),
        ("Hard Gelatin Capsule Shell", "Container", "As required"),
    ]
    for i, h in enumerate(comp_headers):
        comp_tbl.rows[0].cells[i].text = h
    for ri, (ing, fn, amt) in enumerate(comp_data, 1):
        comp_tbl.rows[ri].cells[0].text = ing
        comp_tbl.rows[ri].cells[1].text = fn
        comp_tbl.rows[ri].cells[2].text = amt

    # SmPC section
    doc.add_heading("SUMMARY OF PRODUCT CHARACTERISTICS (SmPC)", level=1)
    doc.add_heading("1. Name of the Medicinal Product", level=2)
    doc.add_paragraph("Amoxicillin 500mg Capsules")
    doc.add_heading("2. Qualitative and Quantitative Composition", level=2)
    doc.add_paragraph(
        "Each capsule contains 500mg of Amoxicillin (as Amoxicillin Trihydrate). "
        "Excipients with known effect: None."
    )
    doc.add_heading("4. Clinical Particulars", level=2)
    doc.add_heading("4.1 Therapeutic Indications", level=3)
    doc.add_paragraph(
        "Amoxicillin 500mg Capsules are indicated for the treatment of the following "
        "infections caused by susceptible organisms: upper respiratory tract infections, "
        "lower respiratory tract infections, urinary tract infections, skin and soft "
        "tissue infections, and dental infections. The prescriber should take into "
        "account official guidance on the appropriate use of antibacterial agents."
    )
    doc.add_heading("4.2 Posology and Method of Administration", level=3)
    doc.add_paragraph(
        "Adults and children over 40kg: 250mg to 500mg three times daily.\n"
        "Children under 40kg: 40mg/kg/day in divided doses.\n"
        "Duration of treatment: 5 to 7 days in most cases.\n"
        "Route: Oral."
    )
    doc.add_heading("4.3 Contraindications", level=3)
    doc.add_paragraph(
        "Hypersensitivity to amoxicillin, any other penicillin, or any of the excipients. "
        "History of severe hypersensitivity reaction (e.g., anaphylaxis) to another "
        "beta-lactam antibiotic (e.g., a cephalosporin, carbapenem, or monobactam)."
    )

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✅ Test dossier saved to: {output_path}")

if __name__ == "__main__":
    make_test_dossier()
